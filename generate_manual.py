"""
Generate the Alpha Move AI User Manual as a formatted Word (.docx) document.
Run: python generate_manual.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0D, 0x1B, 0x3E)   # Deep navy – headings / accents
TEAL        = RGBColor(0x00, 0x8B, 0x8B)   # Dark teal – h2 / table headers
GOLD        = RGBColor(0xD4, 0xAF, 0x37)   # Gold – highlights / callouts
LIGHT_GREY  = RGBColor(0xF4, 0xF6, 0xF8)   # Table row shading
MID_GREY    = RGBColor(0xCC, 0xCC, 0xCC)   # Borders
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x2E)   # Body text
MUTED       = RGBColor(0x55, 0x55, 0x66)   # Captions / notes

# ── Fonts ─────────────────────────────────────────────────────────────────────
FONT_HEADING = "Calibri"
FONT_BODY    = "Calibri"
FONT_MONO    = "Courier New"


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    """Convert RGBColor to 6-char hex string (no #). RGBColor is a tuple."""
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell background with an RGBColor."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = rgb_hex(rgb)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", sz=4):
    """Add thin borders around a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_paragraph_spacing(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "0")
    pPr.append(spacing)


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_break_type())
    remove_paragraph_spacing(para)


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml    import OxmlElement as _OE
    br = _OE("w:br")
    br.set(_qn("w:type"), "page")
    return br   # not used directly – see add_page_break


def add_page_break_v2(doc):
    """Reliable page break via paragraph."""
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    remove_paragraph_spacing(p)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ─────────────────────────────────────────────────────────────────────────────
# Style helpers
# ─────────────────────────────────────────────────────────────────────────────

def para_format(para, space_before=6, space_after=6, line_spacing=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if line_spacing:
        fmt.line_spacing = Pt(line_spacing)


def h1(doc, text):
    """Chapter title – large navy."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name  = FONT_HEADING
    run.font.size  = Pt(26)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    para_format(p, space_before=18, space_after=4)
    # Gold underline bar via border-bottom on the paragraph
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), rgb_hex(GOLD))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(doc, text):
    """Section heading – teal."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = FONT_HEADING
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = TEAL
    para_format(p, space_before=14, space_after=4)
    return p


def h3(doc, text):
    """Sub-section heading – navy."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = FONT_HEADING
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    para_format(p, space_before=10, space_after=3)
    return p


def h4(doc, text):
    """Sub-sub-section – dark teal italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = FONT_HEADING
    run.font.size   = Pt(11)
    run.font.bold   = True
    run.font.italic = True
    run.font.color.rgb = TEAL
    para_format(p, space_before=8, space_after=2)
    return p


def body(doc, text, indent=False):
    """Normal body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = FONT_BODY
    run.font.size  = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    para_format(p, space_before=3, space_after=5, line_spacing=14)
    return p


def body_mixed(doc, parts, indent=False):
    """
    Body paragraph with mixed formatting.
    parts = list of (text, bold, italic, mono, colour)
    colour defaults to DARK_TEXT if None.
    """
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    para_format(p, space_before=3, space_after=5, line_spacing=14)
    for text, bold, italic, mono, colour in parts:
        run = p.add_run(text)
        run.font.name   = FONT_MONO if mono else FONT_BODY
        run.font.size   = Pt(10) if mono else Pt(10.5)
        run.font.bold   = bold
        run.font.italic = italic
        run.font.color.rgb = colour if colour else DARK_TEXT
    return p


def bullet(doc, text, level=0):
    """Bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name  = FONT_BODY
    run.font.size  = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def callout_box(doc, text, label="TIP"):
    """A shaded callout / tip box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, LIGHT_GREY)
    set_cell_borders(cell, color=rgb_hex(TEAL), sz=6)
    cell.width = Inches(6)

    p1 = cell.paragraphs[0]
    run_label = p1.add_run(f"{label}  ")
    run_label.font.name  = FONT_HEADING
    run_label.font.size  = Pt(10)
    run_label.font.bold  = True
    run_label.font.color.rgb = TEAL
    run_body = p1.add_run(text)
    run_body.font.name  = FONT_BODY
    run_body.font.size  = Pt(10)
    run_body.font.color.rgb = DARK_TEXT
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)
    p1.paragraph_format.left_indent  = Cm(0.3)
    p1.paragraph_format.right_indent = Cm(0.3)
    doc.add_paragraph()  # spacer
    return tbl


def disclaimer_box(doc, text):
    """Red-ish disclaimer box."""
    WARN = RGBColor(0xFF, 0xF3, 0xCD)
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, WARN)
    set_cell_borders(cell, color="D4AF37", sz=8)
    p = cell.paragraphs[0]
    r1 = p.add_run("Important Disclaimer  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x7B, 0x5E, 0x07)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.3)
    doc.add_paragraph()


def make_table(doc, headers, rows, col_widths=None, zebra=True):
    """
    Create a nicely formatted table.
    headers : list of str
    rows    : list of list of str
    col_widths : list of floats in cm (optional)
    """
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, hdr in enumerate(headers):
        set_cell_bg(hdr_cells[i], NAVY)
        set_cell_borders(hdr_cells[i], color="FFFFFF", sz=4)
        p   = hdr_cells[i].paragraphs[0]
        run = p.add_run(hdr)
        run.font.name  = FONT_HEADING
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        remove_paragraph_spacing(p)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = LIGHT_GREY if (zebra and r_idx % 2 == 0) else WHITE
        for c_idx, cell_text in enumerate(row_data):
            set_cell_bg(row_cells[c_idx], bg)
            set_cell_borders(row_cells[c_idx], color="CCCCCC", sz=4)
            p   = row_cells[c_idx].paragraphs[0]
            # Bold the first column
            run = p.add_run(str(cell_text))
            run.font.name  = FONT_BODY
            run.font.size  = Pt(9.5)
            run.font.bold  = (c_idx == 0)
            run.font.color.rgb = DARK_TEXT
            remove_paragraph_spacing(p)

    # Column widths
    if col_widths:
        for c_idx, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[c_idx].width = Cm(w)

    doc.add_paragraph()  # spacer after table
    return tbl


def score_badge_para(doc, score_name, score_range, colour):
    """A coloured badge line for score descriptions."""
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0   = tbl.cell(0, 0)
    c1   = tbl.cell(0, 1)
    set_cell_bg(c0, colour)
    c0.width = Cm(4)
    c1.width = Cm(12)
    p0   = c0.paragraphs[0]
    r0   = p0.add_run(f"  {score_name}")
    r0.font.name  = FONT_HEADING
    r0.font.size  = Pt(10)
    r0.font.bold  = True
    r0.font.color.rgb = WHITE
    remove_paragraph_spacing(p0)
    p1   = c1.paragraphs[0]
    r1   = p1.add_run(f"  {score_range}")
    r1.font.name  = FONT_BODY
    r1.font.size  = Pt(10)
    r1.font.color.rgb = DARK_TEXT
    remove_paragraph_spacing(p1)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Big navy banner via a 1-row table spanning full width
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    cell.width = Inches(6.5)

    # Logo at top of banner
    import os
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             "frontend", "public", "logo.png")
    p_logo = cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.4))

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\nAlpha Move AI\nUK Stock Screener\n\n")
    r.font.name  = FONT_HEADING
    r.font.size  = Pt(36)
    r.font.bold  = True
    r.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("User Manual\n")
    r2.font.name  = FONT_HEADING
    r2.font.size  = Pt(20)
    r2.font.bold  = False
    r2.font.color.rgb = GOLD

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("\nAn Investor's Guide to Finding Promising Opportunities\nin UK Markets\n\n")
    r3.font.name  = FONT_BODY
    r3.font.size  = Pt(13)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0xAA, 0xCC, 0xEE)

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("\nEdition: April 2026\n\n")
    r4.font.name  = FONT_BODY
    r4.font.size  = Pt(11)
    r4.font.color.rgb = MID_GREY

    doc.add_paragraph()
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Table of Contents
# ─────────────────────────────────────────────────────────────────────────────

def add_toc(doc):
    h1(doc, "Table of Contents")

    toc_entries = [
        ("1",   "Introduction — What Is a Stock Screener?",            4),
        ("2",   "The Dashboard Layout",                                 5),
        ("3",   "The Screener — Finding Stocks",                        5),
        ("3.1", "  Understanding the Columns",                          5),
        ("3.2", "  The Four Scores Explained",                          6),
        ("3.3", "  Using the Filters",                                  8),
        ("3.4", "  Analyst View",                                       9),
        ("3.5", "  The PEGY Column — Are You Paying a Fair Price?",     9),
        ("3.6", "  The Watchlist — Save Companies to Follow",           10),
        ("3.7", "  Excluding Sectors You Want to Avoid",                10),
        ("4",   "Company Detail — Drilling Into a Stock",               11),
        ("4.1–4.7", "  Chart, Overview, Financials, Valuation, Health, Growth, Analysts", 11),
        ("4.8", "  Company News Tab — Press, RNS & AI Summary",         13),
        ("5",   "Sector Analysis",                                      14),
        ("6",   "Markets",                                              15),
        ("7",   "Analyst Monitor",                                      17),
        ("8",   "RNS News Screener — Catching Catalysts Early",         18),
        ("9",   "Analytics — The Visual Map of the Market",             20),
        ("10",  "The Sidebar — Your Instant Dashboard",                 21),
        ("11",  "How To Find Investment Leads — Step-by-Step Workflows", 22),
        ("12",  "Glossary of Financial Terms",                          26),
        ("A",   "Appendix A — Methodology References",                  28),
        ("B",   "Appendix B — ICB Sector Company List",                 30),
    ]

    for num, title, _ in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r_num   = p.add_run(f"{num}  ")
        r_num.font.bold  = True
        r_num.font.size  = Pt(10.5)
        r_num.font.color.rgb = TEAL if len(num) == 1 else DARK_TEXT
        r_title = p.add_run(title)
        r_title.font.size  = Pt(10.5)
        r_title.font.color.rgb = DARK_TEXT

    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 1 — Introduction
# ─────────────────────────────────────────────────────────────────────────────

def add_section_1(doc):
    h1(doc, "1  Introduction")
    h2(doc, "What Is a Stock Screener?")
    body(doc,
        "Investing in individual shares means choosing from thousands of publicly traded companies. "
        "Without a tool to filter and compare them, the task is overwhelming. A stock screener lets "
        "you set criteria — such as 'show me only companies with low debt and strong profit growth' — "
        "and instantly narrows the universe down to a manageable shortlist."
    )
    body(doc,
        "Alpha Move AI goes further than a basic screener. It calculates composite scores for quality, "
        "momentum, value, and risk; tracks what professional analysts are saying; reads the official "
        "regulatory news feed and uses AI to flag market-moving announcements; and monitors the overall "
        "health of the UK market so you can time your decisions more intelligently."
    )

    h2(doc, "The Investing Goal: Maximum Upside, Minimum Downside")
    body(doc,
        "Every tool in this app is built around two questions every investor needs to answer:"
    )
    bullet(doc, "Upside: What is the realistic chance this share goes up, and by how much?")
    bullet(doc, "Downside: If I am wrong, how badly could I lose?")
    body(doc,
        "A great investment is one with asymmetric returns — the upside is much larger than the "
        "downside. The features in this manual are designed to help you systematically find those "
        "opportunities, by combining four kinds of signals:"
    )
    make_table(doc,
        ["Signal Type", "Tool to use", "What it tells you"],
        [
            ["Quality of the business",  "Quality Score, Piotroski, Health tab",          "Will this company still exist and thrive in 5 years?"],
            ["Price you are paying",     "P/E, P/B, PEGY, Valuation tab",                  "Are you over-paying compared to growth and earnings?"],
            ["Trend & timing",           "Momentum, Sector Rotation, Fear & Greed",        "Is the market currently agreeing with you?"],
            ["Catalysts & sentiment",    "RNS News Screener, Analyst Monitor",             "Is something happening right now that could move the price?"],
        ],
        col_widths=[4, 5, 7]
    )
    callout_box(doc,
        "Workflows in Section 11 show exactly how to combine these signals to find candidates — "
        "including a master 'Maximum Upside / Minimum Downside' workflow (11.8) that stacks every "
        "tool in the app into a single high-conviction filter.",
        label="WHERE TO START"
    )

    disclaimer_box(doc,
        "Nothing in this tool constitutes financial advice. The scores and indicators are research "
        "aids only. Always conduct your own due diligence and consider consulting a qualified "
        "financial adviser before making any investment decisions."
    )

    h2(doc, "The Markets We Cover")
    body(doc,
        "Alpha Move AI focuses on UK-listed equities — companies whose shares trade on the London Stock "
        "Exchange. These are grouped into several indexes:"
    )
    make_table(doc,
        ["Index", "What it contains"],
        [
            ["FTSE 100",       "The 100 largest UK-listed companies by market value (e.g. Shell, AstraZeneca, HSBC). Often called 'the Footsie'."],
            ["FTSE 250",       "The next 250 largest companies. More domestically focused than the FTSE 100."],
            ["FTSE SmallCap",  "Smaller companies below the FTSE 250. Higher growth potential but less liquidity."],
            ["FTSE All-Share", "The combined FTSE 100 + 250 + SmallCap."],
            ["AIM",            "The Alternative Investment Market. Smaller, often younger businesses. Higher risk, higher potential reward."],
        ],
        col_widths=[4, 12]
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 2 — Dashboard Layout
# ─────────────────────────────────────────────────────────────────────────────

def add_section_2(doc):
    h1(doc, "2  The Dashboard Layout")
    body(doc, "When you open Alpha Move AI you see four main areas:")
    for item, desc in [
        ("Top navigation bar",  "Screener, Watchlist, Analysts, RNS News, Analytics, plus a Markets dropdown (Fear & Greed, Cross-Asset, Rotation, Breadth, Signal Log)."),
        ("Left sidebar",        "A live pulse of the market: benchmark returns, market fear levels, and the current cycle signal. Toggle on or off using the icon at the top-left of the navigation bar."),
        ("Main content area",   "Changes depending on which page you are on."),
        ("Search bar",          "Type a company name or ticker symbol to jump straight to its detail page."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{item}:  ")
        r1.font.bold = True
        r1.font.color.rgb = NAVY
        r1.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.color.rgb = DARK_TEXT
        r2.font.size = Pt(10.5)
    callout_box(doc,
        "Direct links to companies: when you open a company page, the URL updates with the ticker "
        "(e.g. …#company/AZN.L). Bookmark or share that URL to land directly on that company. The "
        "browser Back button takes you back to wherever you came from (screener, watchlist, RNS feed).",
        label="TIP"
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 3 — The Screener
# ─────────────────────────────────────────────────────────────────────────────

def add_section_3(doc):
    h1(doc, "3  The Screener — Finding Stocks")
    body(doc,
        "The Screener is the heart of the application. It displays a table of UK stocks which you "
        "can filter and sort to find candidates matching your investment strategy."
    )

    # 3.1
    h2(doc, "3.1  Understanding the Columns")
    h3(doc, "Fundamentals View")
    make_table(doc,
        ["Column", "What it means", "What to look for"],
        [
            ["Symbol",      "The stock's ticker code, e.g. AZN.L. The .L suffix means London-listed.", "—"],
            ["Name",        "Company name.", "—"],
            ["Sector",      "The industry the company operates in (ICB classification).", "Use to compare like-for-like"],
            ["Index",       "Which FTSE index the stock belongs to.", "Indicates company size"],
            ["Mkt Cap",     "Market Capitalisation — total stock market value. Share price × shares outstanding.", "Larger = more stable; smaller = more growth potential"],
            ["P/E",         "Price-to-Earnings — how much investors pay per £1 of profit. A high P/E implies high growth expectations.", "Compare within the same sector"],
            ["P/B",         "Price-to-Book — share price ÷ net asset value per share. P/B < 1 means you are buying £1 of assets for less than £1.", "Useful for banks and property"],
            ["ROE",         "Return on Equity — net profit as % of shareholders' equity. 15%+ is generally considered good.", "Higher and consistent is ideal"],
            ["Rev Growth",  "Year-on-year revenue growth percentage.", "Positive and accelerating"],
            ["D/E",         "Debt-to-Equity — total debt ÷ equity. High D/E increases risk, especially when interest rates are high.", "Lower is safer"],
            ["PEGY",        "P/E ÷ (Growth + Yield). Value-for-money check (Section 3.5).", "Below 1 = great value; above 2 = expensive for growth"],
            ["Momentum",    "Price trend strength score 1–10 (Section 3.2).", "7+ = strong upward trend"],
            ["Quality",     "Consistency and level of returns score 0–10 (Section 3.2).", "7+ = high-quality business"],
            ["Value",       "Piotroski F-Score 0–9 — fundamental health (Section 3.2).", "7+ strong; below 3 = weak"],
            ["Risk",        "Composite risk score 1–10 (Section 3.2).", "Lower = safer"],
            ["★ Star",      "Click the star icon at the start of any row to add or remove the company from your Watchlist (Section 3.6).", "Gold = saved; grey = not saved"],
        ],
        col_widths=[3, 8, 5]
    )

    callout_box(doc,
        "What is Equity?  Think of a company like a house. If the house is worth £300,000 but you "
        "have a £200,000 mortgage, your equity is £100,000. A company's equity is its total assets "
        "minus total liabilities — what shareholders would receive if everything were sold.",
        label="CONCEPT"
    )

    # 3.2
    h2(doc, "3.2  The Four Scores Explained")
    body(doc,
        "These composite scores are calculated from underlying financial data and allow quick "
        "comparison across the full universe without reading each company's accounts individually."
    )

    # Momentum
    h3(doc, "Momentum Score  (1–10)")
    body(doc,
        "Measures how strongly the share price has been trending upward over the medium term. "
        "The score uses the academic concept of 12-1 month price momentum — one of the most "
        "extensively replicated findings in finance."
    )
    h4(doc, "Calculation")
    body(doc,
        "The formula compares the stock's price 63 trading days ago (~3 months) to its price "
        "252 trading days ago (~12 months):"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("Momentum Return = (Price 63 days ago) ÷ (Price 252 days ago) − 1")
    run.font.name = FONT_MONO
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL

    body(doc,
        "The most recent 3 months are deliberately excluded. Research shows very recent returns "
        "tend to reverse (short-term bounce often followed by a pullback), while the 3–12 month "
        "window tends to persist. All stocks are ranked by this return and assigned a score of 1–10, "
        "where 10 = top 10% in the screened universe."
    )
    callout_box(doc,
        "A Momentum Score of 7+ suggests the stock has been among the better performers. "
        "Combined with strong fundamentals, this confirms the market is already recognising "
        "the quality you have identified.  (See Appendix A — Jegadeesh & Titman 1993)",
        label="TIP"
    )

    # Quality
    h3(doc, "Quality Score  (0–10)")
    body(doc,
        "Measures how high and how consistent the company's returns and profit margins are. "
        "Up to 2 points are awarded for each of five criteria, checking both the absolute level "
        "and whether it beats the median of the screened universe:"
    )
    make_table(doc,
        ["Criterion", "Points"],
        [
            ["ROIC > 10% or above median ROIC",              "0–2"],
            ["ROE > 15% or above median ROE",                "0–2"],
            ["Gross Margin > 30% or above median",           "0–2"],
            ["Operating Margin > 10% or above median",       "0–2"],
            ["FCF Margin > 5% or Net Margin above median",   "0–2"],
        ],
        col_widths=[13, 3]
    )
    body(doc,
        "ROIC (Return on Invested Capital) is the best single measure of business quality — a high "
        "ROIC means the company has a competitive advantage (a 'moat') that lets it earn outsized "
        "returns. A Quality Score of 7+ indicates a genuinely high-quality business."
    )

    # Piotroski
    h3(doc, "Piotroski F-Score  (0–9)  —  the 'Value' column")
    body(doc,
        "The Piotroski F-Score measures fundamental financial health and improving momentum. "
        "Originally designed to separate improving value stocks from deteriorating ones, it applies "
        "9 binary tests — each scores 0 or 1:"
    )
    make_table(doc,
        ["Test", "Category", "What it checks"],
        [
            ["1", "Profitability",   "Is Return on Assets (ROA) positive?"],
            ["2", "Profitability",   "Is Operating Cash Flow positive?"],
            ["3", "Profitability",   "Is ROA higher this year than last?"],
            ["4", "Earnings Quality","Is Cash Flow > Net Income? (profit backed by cash, not accounting entries)"],
            ["5", "Leverage",        "Is Debt/Equity lower this year than last?"],
            ["6", "Liquidity",       "Is the Current Ratio higher this year than last?"],
            ["7", "Dilution",        "Has the share count stayed the same or fallen?"],
            ["8", "Efficiency",      "Is Gross Margin higher this year than last?"],
            ["9", "Efficiency",      "Is Asset Turnover higher this year than last?"],
        ],
        col_widths=[1.2, 4, 10.8]
    )
    make_table(doc,
        ["Score", "Interpretation"],
        [
            ["7–9", "Strong — the company is healthy and improving. Classic value investor 'buy' territory."],
            ["4–6", "Mixed — some positives, some concerns. Read the detail before proceeding."],
            ["0–3", "Weak — multiple warning signs. Approach with caution."],
        ],
        col_widths=[3, 13]
    )
    body(doc, "Academic reference: Piotroski (2000) — see Appendix A.")

    # Risk
    h3(doc, "Risk Score  (1–10)")
    body(doc,
        "A composite assessment blending two components: 60% Altman Z-Score and 40% Volatility."
    )
    make_table(doc,
        ["Risk Score", "Interpretation", "Action"],
        [
            ["1–3",  "Low risk — financially robust, low volatility",        "Suitable for all strategies"],
            ["4–5",  "Moderate risk — some leverage or volatility",          "Review health tab before investing"],
            ["6–7",  "Elevated risk — investigate balance sheet carefully",   "Use only with strong quality/value scores"],
            ["8–10", "High risk — distress signals or extreme volatility",    "Avoid unless you understand the specific situation"],
        ],
        col_widths=[3, 7, 6]
    )
    body(doc,
        "The Altman Z-Score was developed in 1968 to predict corporate bankruptcy. "
        "A Z-Score above 3.0 indicates a financially safe company; below 1.8 suggests "
        "significant distress risk. Academic reference: Altman (1968) — see Appendix A."
    )

    # 3.3
    h2(doc, "3.3  Using the Filters")
    h3(doc, "Basic Filters")
    make_table(doc,
        ["Filter", "Practical use"],
        [
            ["Sector",          "Focus on an industry you understand. Click the ⊘ button next to a sector in the dropdown to exclude it (Section 3.7)."],
            ["FTSE Index",      "Filter by company size. FTSE 100 = blue chips; FTSE 250 = mid-cap; FTSE 350 = both; FTSE SmallCap; AIM 100 = largest 100 AIM companies."],
            ["Market Cap (min)","Set a floor to exclude very small, illiquid stocks"],
            ["P/E (max)",       "Setting to 25 focuses on reasonably valued companies; excludes speculative high-PE names"],
            ["ROE (min)",       "Set to 10–15% to screen out unprofitable businesses"],
            ["Rev Growth (min)","Set to 5% to focus on growing companies"],
        ],
        col_widths=[4, 12]
    )
    h3(doc, "Score-Based Filters")
    make_table(doc,
        ["Filter", "Suggested Starting Value", "Rationale"],
        [
            ["Min Momentum",  "6",  "Focus on stocks with above-average price trends"],
            ["Min Quality",   "6",  "Only businesses with strong returns and margins"],
            ["Min Piotroski", "6",  "Financially healthy and improving fundamentals"],
            ["Max Risk",      "5",  "Exclude the higher-risk names from your shortlist"],
        ],
        col_widths=[4, 4, 8]
    )

    # 3.4
    h2(doc, "3.4  Analyst View")
    body(doc,
        "Clicking the Analyst view tab changes the screener columns to show professional analyst data:"
    )
    make_table(doc,
        ["Column", "What it means"],
        [
            ["Consensus",    "The aggregated view of all analysts covering the stock: Buy, Hold, or Sell."],
            ["Upside",       "How far (%) the current price is from the average analyst price target. +20% means analysts see 20% upside."],
            ["Buy%",         "Percentage of covering analysts with a Buy or Strong Buy recommendation."],
            ["# Analysts",   "Number of analysts following the stock. More analysts = more reliable consensus."],
            ["Rev Score",    "Earnings revision score: upward EPS revisions minus downward revisions in past 30 days. Positive = analysts raising forecasts — a strong signal."],
        ],
        col_widths=[3, 13]
    )
    callout_box(doc,
        "Professional analysts spend their careers studying individual sectors and companies. "
        "While no analyst is always right, a strong consensus — especially when combined with "
        "rising earnings estimates — is a meaningful signal worth acting on.",
        label="WHY IT MATTERS"
    )

    # 3.5  PEGY
    h2(doc, "3.5  The PEGY Column — Are You Paying a Fair Price for Growth?")
    body(doc,
        "PEGY stands for Price/Earnings divided by (Growth + Yield). It is one of the simplest and "
        "most useful 'value-for-money' tests for a stock."
    )
    body(doc,
        "The intuition: a high P/E (say 30) sounds expensive — but if the company is growing earnings "
        "at 25% a year and pays a 5% dividend, you are paying 30 ÷ (25 + 5) = a PEGY of 1.0, which is "
        "actually fair. Conversely, a low P/E (say 10) sounds cheap — but if growth is only 2% and the "
        "dividend yield is 1%, the PEGY is 10 ÷ 3 = 3.3, meaning you are paying a lot for very little "
        "growth or income."
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("PEGY = P/E ratio ÷ (Earnings growth % + Dividend yield %)")
    run.font.name = FONT_MONO
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    body(doc,
        "The growth figure is the forward analyst EPS growth (when at least three analysts cover the "
        "stock); otherwise the calculation falls back to the company's 10-year average EPS growth."
    )
    make_table(doc,
        ["PEGY value", "Interpretation", "Colour in screener"],
        [
            ["Below 1",    "Potentially great value — paying less than fair price for the growth and income", "Green"],
            ["1 to 2",     "Fair value",                                                                       "Amber"],
            ["Above 2",    "Expensive relative to growth and income",                                          "Red"],
            ["Blank (—)",  "Not enough data — usually no analyst growth or zero dividend",                     "Grey"],
        ],
        col_widths=[3, 9, 4]
    )
    callout_box(doc,
        "Why PEGY matters for upside / downside: it directly answers 'am I overpaying?'. A low PEGY "
        "combined with a high Quality Score is the textbook setup for asymmetric upside — a solid "
        "business at a fair price, with growth and income working in your favour.",
        label="WHY IT MATTERS"
    )

    # 3.6  Watchlist
    h2(doc, "3.6  The Watchlist — Save Companies to Follow")
    body(doc,
        "The Watchlist is your personal shortlist of companies to follow. To add a company, click the "
        "★ Star icon at the start of any row in the Screener table — the star turns gold and the "
        "company is saved. Click again to remove."
    )
    body(doc,
        "Click 'Watchlist' in the top navigation bar to view your saved companies in the same "
        "screener-style table, with all the same columns and sorting controls. The header shows the "
        "count, e.g. '12 saved'."
    )
    body(doc, "Why use it:")
    bullet(doc, "Build a small focus list of 10–20 companies to monitor closely rather than scanning the whole universe every day.")
    bullet(doc, "Track companies you have flagged for further research after running a screen.")
    bullet(doc, "Keep a 'buy when cheaper' list — companies you believe in but want to wait for a better entry price on.")
    callout_box(doc,
        "The watchlist is stored in your browser. Clearing your browser data resets it. If you use "
        "the app on multiple devices, the watchlists are independent.",
        label="TIP"
    )

    # 3.7  Excluded sectors
    h2(doc, "3.7  Excluding Sectors You Want to Avoid")
    body(doc,
        "Sometimes you want to screen 'all stocks except these sectors' — for example, ethical "
        "preferences (avoid Tobacco), or because you already have heavy exposure to a sector "
        "elsewhere in your portfolio."
    )
    body(doc, "How it works:")
    bullet(doc, "Open the Sector dropdown at the top of the Screener.")
    bullet(doc, "Each sector row has a ⊘ button on the right.")
    bullet(doc, "Click ⊘ to add a sector to your excluded list. Excluded sectors appear with a strikethrough in the dropdown.")
    bullet(doc, "Click ⊘ again to un-exclude.")
    body(doc,
        "Excluded sectors appear as red chips above the screener table (e.g. 'Excluded: Energy ✕'). "
        "Click the ✕ on a chip to remove that exclusion. The dropdown header shows the count, e.g. "
        "'All Sectors (−2)', when exclusions are active."
    )

    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 4 — Company Detail
# ─────────────────────────────────────────────────────────────────────────────

def add_section_4(doc):
    h1(doc, "4  Company Detail — Drilling Into a Stock")
    body(doc,
        "Click any company name to open the Company Detail panel. Seven tabs give you a complete "
        "picture of the business."
    )

    h2(doc, "4.1  Chart Tab")
    body(doc,
        "A 5-year price chart with optional moving average overlays:"
    )
    for label, desc in [
        ("MA20 (20-day moving average)",
         "The average closing price over the last 20 trading days. Used to gauge short-term trend direction. Price above MA20 = short-term uptrend."),
        ("MA50 (50-day moving average)",
         "A medium-term trend indicator. Price above MA50 is broadly considered an uptrend; crossing below is a warning sign."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{label}:  ")
        r1.font.bold = True
        r1.font.color.rgb = TEAL
        r1.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK_TEXT

    h2(doc, "4.2  Overview Tab")
    body(doc, "A snapshot of the most important numbers at a glance:")
    make_table(doc,
        ["Metric", "Plain English", "Good sign"],
        [
            ["Revenue",           "Total sales generated — the 'top line'.",                                        "Growing year on year"],
            ["Net Income",        "Profit after all costs, interest, and tax — the 'bottom line'.",                  "Positive and growing"],
            ["EBITDA",            "Earnings before interest, tax, depreciation & amortisation. A proxy for cash earnings.", "Rising trend"],
            ["FCF",               "Cash left after capital expenditure — available for dividends, buybacks, or debt repayment.", "Positive and growing"],
            ["EPS",               "Earnings per share — profit divided by number of shares.",                        "Rising EPS"],
            ["ROIC",              "Return on Invested Capital — profit for every £1 of capital invested. Best quality metric.", "> 10–15%"],
            ["Current Ratio",     "Current assets ÷ current liabilities. Measures short-term solvency.",             "> 1.0"],
            ["Interest Coverage", "Operating profit ÷ interest expense. How many times the company can pay its interest bill.", "> 5× comfortable; < 2× risky"],
        ],
        col_widths=[3.5, 8, 4.5]
    )

    h2(doc, "4.3  Financials Tab")
    body(doc,
        "Five years of annual income statement data in chart form: Revenue, Gross Profit, "
        "Operating Income, EBITDA, Net Income, and Free Cash Flow."
    )
    callout_box(doc,
        "Look for all five lines trending upward. If revenue grows but net income and FCF do not, "
        "the company may be a 'value trap' — top-line growth is not converting into shareholder value.",
        label="WARNING"
    )

    h2(doc, "4.4  Valuation Tab")
    make_table(doc,
        ["Metric", "What it tells you", "Typical ranges"],
        [
            ["P/E",        "How expensive the stock is relative to earnings.",                 "10–15 cheap; 20–30 fair for growth; 40+ expensive"],
            ["P/B",        "Price relative to book value.",                                   "< 1 may indicate undervalue; > 5 implies high expected returns"],
            ["P/S",        "Useful when a company is not yet profitable.",                    "< 1 cheap; > 5 expensive for most sectors"],
            ["EV/EBITDA",  "Enterprise value relative to operational earnings. Used in buyout valuations.", "8–12 fair; < 6 potentially cheap"],
        ],
        col_widths=[2.5, 8, 5.5]
    )

    h2(doc, "4.5  Health Tab")
    body(doc,
        "Focuses on balance sheet and financial risk. Key items: Current Ratio trend, Net Debt "
        "(negative net debt = more cash than debt = excellent position), Altman Z-Score gauge, "
        "and Debt/Equity trend. A rising D/E ratio combined with falling interest coverage is a "
        "serious warning sign."
    )

    h2(doc, "4.6  Growth Tab")
    make_table(doc,
        ["Metric", "What to look for"],
        [
            ["Revenue CAGR",  "Compound Annual Growth Rate. 10%+ over 5–10 years is strong."],
            ["EPS CAGR",      "Faster than Revenue CAGR = expanding margins."],
            ["FCF CAGR",      "Growing free cash flow is the ultimate sign of a compounding business."],
            ["Margin trends", "Expanding gross, operating, and net margins signal an improving competitive position."],
        ],
        col_widths=[4, 12]
    )

    h2(doc, "4.7  Analysts Tab")
    body(doc,
        "The most detailed analyst view: consensus bar chart, price targets (mean/median/high/low), "
        "EPS and revenue estimates for this year and next, and earnings revision counts."
    )
    callout_box(doc,
        "The revision trend is particularly powerful. When analysts start raising earnings estimates, "
        "it often precedes share price appreciation as the market reprices the improved outlook. "
        "Downward revisions are an equally important early warning signal.",
        label="KEY INSIGHT"
    )

    h2(doc, "4.8  Company News Tab — Press, RNS & AI Summary")
    body(doc,
        "The Company News tab pulls together everything written about the company over the last 6 "
        "months in one place, and lets you ask the AI to summarise it for you."
    )
    body(doc, "There are three sections on this tab:")

    h3(doc, "AI Summary — Last 60 Days")
    body(doc,
        "Press the ✦ Generate summary button at the top of the tab and DeepSeek (an AI model) will "
        "read every regulatory announcement and press article from the last 60 days, then produce: "
        "a short summary paragraph, 2–4 key themes (e.g. 'Margin pressure from input costs', "
        "'Successful product launch in Asia'), and a 'Watch Next' note pointing to the events that "
        "will determine whether the story continues or breaks. Click ↻ Regenerate when new news has come in."
    )

    h3(doc, "Regulatory (RNS) Feed")
    body(doc,
        "The official Stock Exchange announcements the company has filed (results, M&A, director "
        "dealings, contract wins, capital raises, etc.). Each row shows:"
    )
    bullet(doc, "Tier badge (A / B / C) — the AI's view of how significant the announcement is. See Section 8.")
    bullet(doc, "Headline — click to open the full announcement on Investegate.")
    bullet(doc, "AI thesis — a one-sentence interpretation of why the announcement matters.")
    bullet(doc, "Action pill — BUY / WATCH / AVOID / NEUTRAL — the AI's suggested response.")

    h3(doc, "Press / Google News Feed")
    body(doc,
        "Articles from the wider press (FT, Reuters, Bloomberg, broker notes covered in the press) "
        "sourced from Google News. Cached for 24 hours per company; click ↻ Refresh news to force "
        "a fresh fetch."
    )

    callout_box(doc,
        "Many of the strongest one-week share price moves happen on RNS announcements. Reading the "
        "recent RNS history tells you what catalysts have already played out — and helps you "
        "anticipate what might come next. The AI summary is especially useful when first researching "
        "a name: 60 days of news condensed into a few paragraphs you can read in 30 seconds.",
        label="WHY IT MATTERS"
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 5 — Sector Analysis
# ─────────────────────────────────────────────────────────────────────────────

def add_section_5(doc):
    h1(doc, "5  Sector Analysis")

    h2(doc, "5.1  Sector Rotation")
    body(doc,
        "Different sectors of the economy perform better at different points in the economic cycle. "
        "Understanding which phase we are in can help you tilt your portfolio toward areas likely "
        "to outperform."
    )
    make_table(doc,
        ["Cycle Phase", "Leading Sectors", "Characteristics"],
        [
            ["Recovery",    "Industrials, Consumer Discretionary, Financials", "Economy picking up after a downturn; earnings start to recover"],
            ["Expansion",   "Technology, Materials, Energy",                   "Strong growth; risk appetite high; commodity demand rises"],
            ["Slowdown",    "Consumer Staples, Health Care",                   "Growth moderating; investors rotate to defensives"],
            ["Contraction", "Utilities, Consumer Staples, Gilts",              "Recession; capital preservation priorities; defensive havens"],
        ],
        col_widths=[3.5, 6, 6.5]
    )
    h3(doc, "The Relative Strength (RS) Score")
    body(doc, "For each sector, Alpha Move AI calculates:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("RS Score = (Sector 63-day return) ÷ (FTSE All-Share 63-day return)")
    run.font.name = FONT_MONO
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    make_table(doc,
        ["RS Score", "Signal", "Meaning"],
        [
            ["> 1.05",        "BUY",     "Sector outperformed the market by > 5% over the past 3 months — leadership signal"],
            ["0.95 – 1.05",   "NEUTRAL", "In line with the market"],
            ["< 0.95",        "AVOID",   "Sector underperformed by > 5% — laggard signal"],
        ],
        col_widths=[3, 3, 10]
    )

    h2(doc, "5.2  Market Breadth")
    make_table(doc,
        ["Indicator", "Healthy reading", "Warning reading"],
        [
            ["% stocks above 50-day MA", "> 70% — broad participation", "< 40% — narrow market, index carried by few stocks"],
            ["52-Week Highs vs Lows",    "Far more highs than lows",    "More lows than highs — deteriorating market"],
            ["Advance/Decline Line",     "Rising — confirms index move", "Diverging (A/D falls while index rises) — hidden weakness"],
        ],
        col_widths=[5, 5, 6]
    )

    h2(doc, "5.3  Signal Log")
    body(doc,
        "A chronological record of automatically generated signals: BUY (sector RS crossed above 1.05), "
        "AVOID (RS dropped below 0.95), ALERT (breadth changes), and INFO. Use it to track how long "
        "a signal has been in place and evaluate it with hindsight."
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 6 — Markets
# ─────────────────────────────────────────────────────────────────────────────

def add_section_6(doc):
    h1(doc, "6  Markets")

    h2(doc, "6.1  Fear & Greed Index")
    body(doc,
        "Alpha Move AI's UK Fear & Greed Index combines six independent data points into a single score "
        "from 0 (Extreme Fear) to 100 (Extreme Greed), purpose-built for UK markets."
    )
    callout_box(doc,
        "Markets are driven by two emotions: fear and greed. When greed dominates (score near 100), "
        "assets are often overpriced and a correction may be coming. When fear dominates (score near 0), "
        "assets are often underpriced and recovery may be near. As Warren Buffett said: "
        "'Be fearful when others are greedy, and greedy when others are fearful.'",
        label="PHILOSOPHY"
    )
    make_table(doc,
        ["Component", "What it measures"],
        [
            ["FTSE Momentum",       "How far the FTSE 100 is above/below its 125-day MA, statistically normalised."],
            ["Market Breadth",      "% of FTSE stocks above their 50-day MA. Broad participation = greed; narrow = fear."],
            ["VIX",                 "Expected US market volatility (inverted — high VIX = low/fear score)."],
            ["Safe Haven Demand",   "Spread between FTSE 100 and UK Gilt ETF 20-day returns. Investors fleeing to bonds = fear."],
            ["Realised Volatility", "Actual 20-day volatility of the FTSE 100. High volatility = fear."],
            ["New Highs vs Lows",   "Ratio of stocks hitting 52-week highs versus lows. More highs = greed."],
        ],
        col_widths=[4.5, 11.5]
    )
    make_table(doc,
        ["Score", "Label", "Implication for investors"],
        [
            ["75–100", "Extreme Greed", "Market may be overheated. Consider reducing risk or locking in profits."],
            ["55–74",  "Greed",         "Positive momentum. Maintain positions but stay selective."],
            ["45–54",  "Neutral",       "No strong directional signal."],
            ["25–44",  "Fear",          "Market under stress. Look for quality stocks on sale."],
            ["0–24",   "Extreme Fear",  "Potential buying opportunity for long-term, patient investors."],
        ],
        col_widths=[2, 3.5, 10.5]
    )

    h2(doc, "6.2  Cross-Asset Monitor")
    make_table(doc,
        ["Asset", "Relevance to UK investors"],
        [
            ["GBP/USD",               "Sterling strength. Rising pound is positive for importers but hurts FTSE 100 exporters (overseas earnings worth less in GBP)."],
            ["Brent Crude Oil",       "Affects Energy sector directly; indirectly affects inflation and consumer spending."],
            ["Gold",                  "Rising gold prices typically indicate global uncertainty or inflation fears — a 'risk-off' signal."],
            ["Gilt Yields vs Utilities", "When gilt yields rise sharply, investors can earn more from government bonds risk-free, reducing the relative appeal of high-dividend sectors like Utilities."],
        ],
        col_widths=[4, 12]
    )

    h2(doc, "6.3  Gilt Yields")
    body(doc,
        "UK government bonds (gilts) are loans to the UK government. The yield is the effective "
        "interest rate paid. Gilt yields are the foundation of all UK asset pricing."
    )
    make_table(doc,
        ["Yield Curve Shape", "What it signals"],
        [
            ["Normal (long rates > short rates)",   "Economic optimism; investors expect growth and some inflation."],
            ["Flat (rates similar across maturities)", "Uncertainty; transition between cycle phases."],
            ["Inverted (short rates > long rates)", "A historically reliable recession warning signal. Bond investors expect rates to fall as the economy weakens."],
        ],
        col_widths=[5.5, 10.5]
    )
    body(doc,
        "Rising gilt yields increase the 'risk-free rate' — the return investors can earn without "
        "any equity risk — which mechanically reduces the relative attractiveness of equities, "
        "particularly long-duration growth stocks."
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 7 — Analyst Monitor
# ─────────────────────────────────────────────────────────────────────────────

def add_section_7(doc):
    h1(doc, "7  Analyst Monitor")
    body(doc,
        "The Analyst Monitor provides a dedicated view of professional analyst sentiment across "
        "all covered stocks."
    )
    for section, desc in [
        ("Latest Consensus",  "All stocks with analyst coverage, sortable by consensus rating, upside %, buy %, or revision score."),
        ("Changes",           "Flags stocks where consensus or upside estimate has shifted materially — often the most actionable signals, as a change in analyst view can precede significant price moves."),
        ("Composite Score",   "A single ranking combining: Buy% (full weight) + Upside% × 0.5 (capped at 50%) + Revision Score × 10. Revision momentum is given extra weight as it is forward-looking."),
        ("Shrinkage adjustment", "When only 1–2 analysts cover a stock, a '100% bullish' rating is usually noise rather than signal. The composite score therefore pulls thinly-covered stocks back toward a neutral 50% baseline — a single analyst at 100% Buy counts as ~58%, while a 20-analyst stock at 80% Buy stays close to its raw value. This deliberately rewards stocks where many independent analysts agree."),
        ("Top Bullish / Bearish", "The five stocks with the strongest positive and negative composite scores — your instant professional-consensus shortlist."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run(f"{section}:  ")
        r1.font.bold = True
        r1.font.color.rgb = TEAL
        r1.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK_TEXT
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 8 — RNS News Screener
# ─────────────────────────────────────────────────────────────────────────────

def add_section_8(doc):
    h1(doc, "8  RNS News Screener — Catching Catalysts Early")
    body(doc,
        "The RNS News Screener (top nav → 'RNS News') is one of the most powerful tools in the app "
        "for catching catalysts early."
    )

    h2(doc, "8.1  What is RNS?")
    body(doc,
        "RNS stands for Regulatory News Service — the official Stock Exchange channel that all listed "
        "UK companies must use to release price-sensitive information. Examples include results "
        "announcements, profit warnings, mergers and offer announcements, major contract wins, drug "
        "approvals, drilling results, director share dealings, and capital raises."
    )
    body(doc,
        "When a company files an RNS, the market reacts quickly — often within seconds for large items. "
        "Reading and reacting to RNS news is one of the most direct sources of edge for retail "
        "investors, because the same announcement reaches you at exactly the same moment as everyone "
        "else. The challenge is volume: there are hundreds of RNS releases every trading day, most of "
        "which are administrative (changes in share count, voting rights, etc.). The RNS News Screener "
        "solves this with an AI pipeline that filters and ranks announcements automatically."
    )

    h2(doc, "8.2  The Two-Layer AI Pipeline")
    body(doc,
        "When you press ↻ Refresh + AI rank, the app runs three stages in sequence (you can watch the "
        "live stage indicator on the button):"
    )
    make_table(doc,
        ["Stage", "What it does"],
        [
            ["Ingest",     "Pulls every new RNS announcement from Investegate (the standard data source for UK regulatory news)."],
            ["Summaries",  "Fetches the AI-generated summary for each announcement from Investegate."],
            ["Rank",       "Sends each Tier A and Tier B item to DeepSeek (an AI model) for ranking, scoring, and a 'what to do about it' recommendation."],
        ],
        col_widths=[3, 13]
    )
    body(doc,
        "The pipeline takes 1–3 minutes per refresh. It also runs automatically several times per UK "
        "trading day (every 15 minutes during the morning RNS window), so the feed is usually fresh."
    )

    h3(doc, "The Tier System")
    body(doc,
        "A simple keyword-based rules classifier does a first-pass coarse sort:"
    )
    make_table(doc,
        ["Tier", "Label", "Examples"],
        [
            ["A", "Significant",  "Profit warnings, full-year results, firm offers (Rule 2.7), strategic reviews, drug approvals, major contract wins"],
            ["B", "Noteworthy",   "Trading updates, possible offers (Rule 2.4), capital raises, drill results, board changes, dividend changes"],
            ["C", "Routine",      "Total voting rights, holdings notifications, PDMR transactions, AGM admin"],
        ],
        col_widths=[1, 3, 12]
    )
    body(doc,
        "Tier A and Tier B items are sent to the AI for ranking. Tier C items are kept in the "
        "database but not ranked (you can still see them by setting Min Score = 0 and Tier filter = C)."
    )

    h2(doc, "8.3  Reading the Feed")
    body(doc,
        "Summary cards at the top of the page show the count of Tier A (orange), Tier B (blue), AI-Ranked "
        "items (green), and Total in feed for your current filter window."
    )
    h3(doc, "Controls")
    make_table(doc,
        ["Control", "What it does"],
        [
            ["Window",         "Look back 6 hours up to 1 week."],
            ["Min score",      "Hide items below this rules-classifier score (e.g. 60 = Tier A+ only)."],
            ["Tier pills",     "Filter by Tier A, Tier B, Tier C, or All."],
            ["Sort: AI / Time", "Either rank by the AI's score (best first) or chronological (newest first)."],
            ["Search box",     "Free-text filter by ticker, company name, or headline keyword."],
        ],
        col_widths=[3.5, 12.5]
    )
    h3(doc, "Each Row Shows")
    bullet(doc, "Time — when the announcement was filed (e.g. 08:23 today, 23 Apr 14:47 earlier).")
    bullet(doc, "Tier badge — A / B / C.")
    bullet(doc, "Ticker / Company — click the ticker to open the company detail page.")
    bullet(doc, "Headline — click to open the full announcement on Investegate.")
    bullet(doc, "AI thesis — a one-sentence interpretation of why the announcement matters.")
    bullet(doc, "AI risks — what could go wrong with the bullish reading.")
    bullet(doc, "Category — the rules-classifier category (Profit Warning, Trading Update, etc.).")
    bullet(doc, "Rules score (0–100) — the rules-classifier importance.")
    bullet(doc, "AI score (0–100) — the DeepSeek score, factoring in valuation, analyst views, and recent price action.")
    bullet(doc, "Action pill — BUY / WATCH / AVOID / NEUTRAL.")

    h2(doc, "8.4  Action Pills — BUY / WATCH / AVOID")
    body(doc,
        "The AI assigns each ranked announcement an action, which is the simplest way to scan the "
        "feed quickly:"
    )
    make_table(doc,
        ["Action", "Colour", "What it suggests"],
        [
            ["BUY",     "Green", "A clear bullish catalyst with limited downside risk in the AI's reading. Worth investigating today."],
            ["WATCH",   "Amber", "A mixed or interesting signal — not a clear buy yet, but worth tracking follow-up news."],
            ["AVOID",   "Red",   "A clear negative — profit warning, going concern, etc. The AI sees more downside than upside."],
            ["NEUTRAL", "Grey",  "Material announcement but no clear directional signal."],
        ],
        col_widths=[2.5, 2.5, 11]
    )
    callout_box(doc,
        "The action is the AI's suggestion based on the announcement plus the company's valuation, "
        "analyst consensus, and recent price context. It is not a buy/sell recommendation — always "
        "cross-check by opening the company detail page and reading the full RNS yourself before acting.",
        label="IMPORTANT"
    )

    h3(doc, "Using the RNS Feed for Upside Hunting")
    body(doc, "The clearest opportunities tend to come from:")
    bullet(doc, "Tier A 'BUY' items where the company also has a high Quality Score and reasonable PEGY — quality + catalyst + fair price = highest-conviction setup.")
    bullet(doc, "Multiple Tier B 'WATCH' items in the same week on the same company — when several positive operational updates accumulate, the market often hasn't fully priced in the cumulative effect yet.")
    bullet(doc, "Recommended offer (Rule 2.7) announcements — these instantly mark the share price near the offered price, but sometimes there is a small 'merger arbitrage' gap if you trust the deal will close.")

    h3(doc, "Using the RNS Feed for Downside Protection")
    bullet(doc, "Profit warnings are listed as Tier A AVOID. If you hold the stock and see one, your downside is real and immediate — read the announcement immediately.")
    bullet(doc, "Going concern statements, suspensions, and delistings are tagged in the Category column — these are some of the worst things that can happen to a stock you own.")
    bullet(doc, "Strategic Reviews can go either way — sometimes leading to break-up value (upside), sometimes to forced asset sales (downside). Read the AI thesis carefully.")

    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 9 — Analytics (Visual Map)
# ─────────────────────────────────────────────────────────────────────────────

def add_section_9_analytics(doc):
    h1(doc, "9  Analytics — The Visual Map of the Market")
    body(doc,
        "The Analytics page (top nav → 'Analytics') shows the entire UK universe as a scatter plot, "
        "so you can spot opportunities visually. Each dot is one company; dot size scales with "
        "market capitalisation."
    )

    h2(doc, "Mode 1 — Quality × PEGY")
    body(doc,
        "X axis: PEGY (lower = cheaper for the growth and income on offer). "
        "Y axis: Quality Score (higher = better business)."
    )
    make_table(doc,
        ["Quadrant", "Meaning", "Colour"],
        [
            ["Top-left",     "High Quality, Low PEGY — Cheap quality (the Buffett quadrant)",   "Green dots"],
            ["Top-right",    "High Quality, High PEGY — Expensive quality",                     "Amber dots"],
            ["Bottom-left",  "Low Quality, Low PEGY — Cheap low-quality (potential value trap)", "Amber dots"],
            ["Bottom-right", "Low Quality, High PEGY — Avoid",                                  "Red dots"],
        ],
        col_widths=[3, 9, 4]
    )
    callout_box(doc,
        "Most asymmetric upside lives in the green top-left quadrant — high-quality businesses "
        "trading at fair-or-below prices for the growth they are delivering. This is the cluster "
        "to mine when hunting for long-term winners.",
        label="WHERE TO HUNT"
    )

    h2(doc, "Mode 2 — Momentum × Risk")
    body(doc,
        "X axis: Risk Score (lower = safer, displayed left-to-right with the safer side on the left). "
        "Y axis: Momentum Score (higher = stronger price trend)."
    )
    make_table(doc,
        ["Quadrant", "Meaning", "Colour"],
        [
            ["Top-left",     "Strong momentum + safe",       "Green dots"],
            ["Top-right",    "Strong momentum + risky",      "Amber dots"],
            ["Bottom-left",  "Weak momentum + safe",         "Amber dots"],
            ["Bottom-right", "Weak momentum + risky",        "Red dots"],
        ],
        col_widths=[3, 9, 4]
    )
    callout_box(doc,
        "Top-left is the goal: a stock trending up and with a low risk score has the best "
        "risk-adjusted profile in the universe. Top-right (high momentum but risky) is the "
        "speculative quadrant — can pay off but blow-ups are sudden.",
        label="WHERE TO HUNT"
    )

    h2(doc, "Controls")
    bullet(doc, "Index pills — limit the plot to FTSE 100, FTSE 250, FTSE 350, or All.")
    bullet(doc, "X-axis zoom slider — pull leftwards to zoom into the dense central area; the indicator at the right tells you how many points are off-chart.")
    bullet(doc, "Click any dot to open that company's detail page directly.")

    callout_box(doc,
        "A scatter plot of 1,000 stocks shows patterns the screener table cannot — clusters of "
        "similar companies, sector concentrations, and isolated outliers (often the most "
        "interesting opportunities). Use it to find ideas you would not have thought to filter for.",
        label="WHY IT MATTERS"
    )

    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 10 — Sidebar (was Section 8)
# ─────────────────────────────────────────────────────────────────────────────

def add_section_10_sidebar(doc):
    h1(doc, "10  The Sidebar — Your Instant Dashboard")
    body(doc, "The sidebar is always visible (toggle on/off via the navigation icon) and gives you an immediate market pulse:")
    make_table(doc,
        ["Item", "What it shows"],
        [
            ["FTSE 100 / 250 / All-Share", "Total return for each index (price + dividends reinvested)."],
            ["VIX",                        "Current US volatility index. Below 15 = calm; 20–30 = elevated; 30+ = stressed markets."],
            ["CNN Fear & Greed",            "External US market sentiment gauge (sourced from CNN)."],
            ["UK Fear & Greed",             "Alpha Move AI's own UK-specific sentiment score (0–100)."],
            ["Cycle Phase",                 "Current estimated market cycle phase (Recovery / Expansion / Slowdown / Contraction)."],
            ["Top RS Sector",               "The sector with the highest relative strength — where market leadership currently sits."],
            ["Market Breadth",              "Overall % of FTSE stocks above their 50-day moving average."],
        ],
        col_widths=[4.5, 11.5]
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 11 — Investment Workflows (was Section 9)
# ─────────────────────────────────────────────────────────────────────────────

def add_section_11_workflows(doc):
    h1(doc, "11  How To Find Investment Leads")
    body(doc,
        "The following workflows use combinations of the tools described above to identify "
        "promising investment candidates. These are starting points for research, not buy "
        "recommendations."
    )
    body(doc,
        "The first five (11.1–11.5) are classic single-angle workflows. The last three (11.6–11.8) "
        "combine multiple tools in the app and are designed specifically to maximise the chance of "
        "upside while minimising downside risk."
    )

    workflows = [
        {
            "num": "11.1",
            "title": "The Quality + Value Filter",
            "suitable": "Long-term investors seeking high-quality businesses at reasonable prices.",
            "philosophy": (
                "The best investments combine high business quality (durable competitive advantages, "
                "high returns on capital) with a reasonable price. This combines elements of Warren "
                "Buffett's quality-focused approach with Piotroski's systematic value screening."
            ),
            "steps": [
                "Go to the Screener (Fundamentals view).",
                "Set filters: Min Quality Score 7, Min Piotroski 6, Max Risk Score 5, Min ROE 12%.",
                "Sort by Quality Score descending.",
                "For each result, open the Valuation tab and assess P/E and EV/EBITDA relative to the sector average.",
                "Check the Growth tab to confirm margins are stable or expanding.",
                "Check the Health tab to confirm the balance sheet is robust.",
            ],
            "look_for": (
                "Companies with Quality 8+, Piotroski 7+, trading on a P/E below their historical "
                "average or sector peers. These are businesses the market may be undervaluing despite "
                "their fundamental strength."
            ),
        },
        {
            "num": "11.2",
            "title": "The Momentum Trend-Following Approach",
            "suitable": "Investors comfortable with medium-term active management.",
            "philosophy": (
                "Stocks with strong price momentum tend to continue outperforming over the following "
                "3–12 months — one of the most replicated findings in academic finance. Combining "
                "momentum with a quality check reduces the risk of buying fundamentally weak "
                "'crowded trades'."
            ),
            "steps": [
                "Check the Fear & Greed Index — only use this strategy when the score is 45–80 (Neutral to Greed). In Extreme Fear, momentum strategies frequently fail.",
                "Check Sector Rotation for the top 2–3 sectors by RS Score (BUY-signalled sectors).",
                "In the Screener, filter by those leading sectors.",
                "Set: Min Momentum Score 7, Min Quality Score 5, Max Risk Score 6.",
                "Sort by Momentum Score descending.",
                "For each result, confirm on the Chart tab that the price is above both MA20 and MA50.",
            ],
            "look_for": (
                "High-momentum stocks in leading sectors with the price above both moving averages. "
                "This combination — sector leadership + stock momentum + price confirmation — "
                "increases the probability of continued outperformance."
            ),
        },
        {
            "num": "11.3",
            "title": "The Analyst Upgrade Hunt",
            "suitable": "Investors who want to follow professional money.",
            "philosophy": (
                "Analyst earnings upgrades (upward EPS revisions) are one of the strongest "
                "short-to-medium term price catalysts. When analysts raise profit forecasts, it "
                "signals they have discovered something positive, and the market typically reprices."
            ),
            "steps": [
                "Go to the Analyst Monitor page.",
                "Sort the Latest Consensus table by Rev Score descending.",
                "Look for stocks with: positive Rev Score AND Buy consensus AND Upside % ≥ 15%.",
                "Cross-reference these names in the Screener to confirm Quality ≥ 6 and Risk Score ≤ 6.",
                "Check the Analysts tab in Company Detail for the individual breakdown of estimate changes.",
            ],
            "look_for": (
                "Stocks where analysts are becoming increasingly positive AND the fundamentals "
                "support the bullish case. The Changes section highlights stocks where consensus "
                "has recently shifted — the freshest signals."
            ),
        },
        {
            "num": "11.4",
            "title": "The Sector Rotation Strategy",
            "suitable": "Investors who want to tilt their portfolio toward market-leading areas.",
            "philosophy": (
                "Markets rotate through sectors as the economic cycle evolves. Being in the leading "
                "sectors significantly improves returns; avoiding lagging sectors reduces drawdowns."
            ),
            "steps": [
                "Go to Sector Analysis → Rotation.",
                "Identify sectors with a BUY signal (RS > 1.05, rising trend) and strong breadth (> 60% of stocks above 50-day MA).",
                "Check the Cycle Wheel — does the suggested phase align with your macro view?",
                "In the Screener, filter by the leading sector(s).",
                "Apply: Min Momentum 6, Min Quality 5. Sort by Market Cap descending for the most liquid names.",
                "Review the Signal Log for context on how long the signal has been in place.",
            ],
            "look_for": (
                "The sweet spot is a sector that has recently generated a BUY signal, not one that "
                "has been leading for 12 months and may be due to rotate. Combine with stock "
                "quality and momentum scores to find the best names within the sector."
            ),
        },
        {
            "num": "11.5",
            "title": "The Defensive Screen (Capital Preservation)",
            "suitable": "Investors prioritising protecting capital in uncertain markets.",
            "philosophy": (
                "When Fear & Greed enters Fear territory, focus shifts to capital preservation. "
                "Defensive stocks — Consumer Staples, Health Care, Utilities — hold their value "
                "better in downturns because demand for their products is relatively stable "
                "regardless of the economic cycle."
            ),
            "steps": [
                "Check the Fear & Greed Index — if score < 40, the defensive screen is most relevant.",
                "Filter by Sector: Consumer Staples, Health Care, or Utilities.",
                "Set: Max Risk Score 4, Min Quality Score 6, Min Piotroski 5.",
                "In the Health tab, look for companies with Net Debt negative (cash-rich) or Interest Coverage > 5×.",
                "In the Analysts tab, confirm consensus is at least Hold — avoid Sell-rated stocks in a declining market.",
            ],
            "look_for": (
                "High-quality, financially robust companies in defensive sectors. In a bear market, "
                "these companies decline less and recover faster — preserving your capital to "
                "redeploy when Fear & Greed signals a recovery."
            ),
        },
        {
            "num": "11.6",
            "title": "The Catalyst Hunt — RNS-Driven Upside",
            "suitable": "Investors who can act on the same day news arrives.",
            "philosophy": (
                "Most large positive share-price moves happen on a specific day in response to a "
                "specific announcement. The RNS News Screener is designed to catch these as they "
                "happen, rank them by significance, and tell you which ones the AI sees as bullish "
                "opportunities versus risk events."
            ),
            "steps": [
                "Open the RNS News page first thing in the morning UK time (most RNS releases come between 07:00 and 08:00).",
                "Set Window to 24 hours, Min score to 60 (Tier A+), and Sort: AI score.",
                "Look for items with the BUY action pill in green, especially those with an AI score ≥ 75.",
                "For each candidate, click the ticker to open the company page.",
                "On the company page check three things: Quality Score ≥ 6, Risk Score ≤ 6, and Analysts tab Upside % > 10.",
                "If all three pass, read the full RNS by clicking the headline — make sure the AI thesis matches what the company actually said.",
            ],
            "look_for": (
                "A high-quality, low-risk company with an AI-flagged BUY catalyst that has not yet "
                "been fully priced by analysts."
            ),
        },
        {
            "num": "11.7",
            "title": "The Visual Map Workflow — Cheap Quality Hunting",
            "suitable": "Long-term investors who want to find high-quality businesses at fair-or-better prices.",
            "philosophy": (
                "Plotting the entire universe on a Quality vs PEGY chart instantly reveals where the "
                "asymmetric opportunities are. The top-left 'cheap quality' quadrant is where "
                "Buffett-style 'wonderful businesses at a fair price' cluster."
            ),
            "steps": [
                "Go to Analytics. Select mode Quality × PEGY.",
                "Filter to FTSE 350 (the largest 350 companies — best liquidity).",
                "Use the X-zoom slider to zoom into the dense region (PEGY 0–3).",
                "Look for green dots in the top-left quadrant — high Quality, low PEGY.",
                "Click each green dot to open the company. Run standard checks: Health (balance sheet), Growth (revenue/profit/FCF rising), Analysts (consensus agrees), Company News (no negative catalysts).",
                "Switch to Momentum × Risk. A company in the top-left of both charts is a rare and powerful setup.",
            ],
            "look_for": (
                "Companies that appear in the green quadrant of both visualisations — the rare "
                "intersection of 'good business, fairly priced, currently trending up, low risk'."
            ),
        },
        {
            "num": "11.8",
            "title": "The Combined 'Maximum Upside / Minimum Downside' Workflow",
            "suitable": "Any investor who wants a systematic way to assemble a high-conviction shortlist.",
            "philosophy": (
                "No single signal is reliable on its own. The strongest investment cases stack "
                "multiple independent signals — quality, value, momentum, professional consensus, "
                "and a recent catalyst — that all point the same way. When five independent things "
                "all say 'buy', the probability of being wrong falls dramatically. This is how you "
                "maximise upside while minimising downside."
            ),
            "steps": [
                "Macro check: open the Sidebar. UK Fear & Greed score 25–55 = excellent hunting; 55–75 = hunt selectively; 75+ = be defensive (use Workflow 11.5 instead). Note the Cycle Phase.",
                "Sector battlefield: Markets → Rotation. Note the top 3 sectors by RS Score with rising trend; avoid sectors with RS < 0.95.",
                "Narrow the universe in the Screener: FTSE 350, leading sectors only, Min Quality ≥ 7, Min Piotroski ≥ 6, Max Risk ≤ 5, Min Momentum ≥ 6, Min Upside % ≥ 10, Consensus = Buy.",
                "Visual cross-check on Analytics: confirm survivors land in the green top-left of both Quality × PEGY and Momentum × Risk modes.",
                "Catalyst check: for each survivor, open Company Detail → News tab. Press ✦ Generate AI summary. Look for positive Tier A/B items in last 30 days. Drop names with profit warnings or going-concern flags.",
                "Final downside checks: Health tab (Net Debt manageable, Z-Score not in distress), Valuation (P/E in line with sector), Growth (revenue/EPS/FCF rising), Analysts (Rev Score positive).",
                "Save and monitor: star surviving names with the ★ icon — they go to your Watchlist for daily review. Re-run the macro check before each buy decision.",
            ],
            "look_for": (
                "A small portfolio (5–15 names) where eight or nine of these signals are GREEN: "
                "macro backdrop, sector leadership, business quality, valuation, trend, safety, "
                "professional consensus, catalysts, and visual confirmation. When fewer than five "
                "are green, walk away — there will be better opportunities."
            ),
        },
    ]

    for wf in workflows:
        h2(doc, f"{wf['num']}  {wf['title']}")
        # Suitable for / Philosophy
        for label, key in [("Suitable for", "suitable"), ("Philosophy", "philosophy")]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(4)
            r1 = p.add_run(f"{label}:  ")
            r1.font.bold = True
            r1.font.color.rgb = NAVY
            r1.font.size = Pt(10.5)
            r2 = p.add_run(wf[key])
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = DARK_TEXT

        h4(doc, "Step-by-step")
        for i, step in enumerate(wf["steps"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r1 = p.add_run(f"  {i}.  ")
            r1.font.bold = True
            r1.font.color.rgb = TEAL
            r1.font.size = Pt(10.5)
            r2 = p.add_run(step)
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = DARK_TEXT

        callout_box(doc, wf["look_for"], label="WHAT TO LOOK FOR")

    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 12 — Glossary (was Section 10)
# ─────────────────────────────────────────────────────────────────────────────

def add_section_12_glossary(doc):
    h1(doc, "12  Glossary of Financial Terms")
    glossary = [
        ("Annual Report",       "A yearly document published by a listed company reporting its financial results, strategy, and governance."),
        ("Bear Market",         "A sustained market decline of 20% or more from recent highs."),
        ("Bull Market",         "A sustained market rise of 20% or more from recent lows."),
        ("CAGR",                "Compound Annual Growth Rate — the smoothed annual rate of growth over a multi-year period."),
        ("Capital Gain",        "Profit made when a share price rises above the price you paid."),
        ("Current Ratio",       "Current assets ÷ current liabilities. Above 1.0 means short-term bills can be met."),
        ("Dividend",            "A cash payment made by a company to its shareholders, usually quarterly or annually."),
        ("Dividend Yield",      "Annual dividend per share ÷ share price. Expressed as a percentage."),
        ("EPS",                 "Earnings Per Share — company net profit divided by shares outstanding."),
        ("Enterprise Value",    "Market cap + net debt. The theoretical takeover cost of a business."),
        ("Equity",              "Shareholders' ownership stake. Total assets minus total liabilities."),
        ("FCF",                 "Free Cash Flow — operating cash flow minus capital expenditure."),
        ("Gilts",               "UK government bonds. The 'risk-free' investment benchmark for UK investors."),
        ("Gross Margin",        "(Revenue − Cost of Goods Sold) ÷ Revenue. Profitability before operating costs."),
        ("ICB",                 "Industry Classification Benchmark — the standard sector taxonomy for UK-listed companies."),
        ("Liquidity",           "How easily a share can be bought or sold without moving its price."),
        ("Market Cap",          "Share price × shares outstanding. The total market value of the company."),
        ("Moving Average",      "The average price over a specified number of recent days. Smooths daily price noise."),
        ("Net Debt",            "Total borrowings minus cash. Negative net debt = more cash than debt."),
        ("P/B Ratio",           "Price to Book — share price divided by net assets per share."),
        ("P/E Ratio",           "Price to Earnings — share price divided by earnings per share."),
        ("PEGY",                "P/E divided by (Growth + Yield). A value-for-money check that combines the P/E with both growth and dividend yield in one figure."),
        ("P/S Ratio",           "Price to Sales — market cap divided by annual revenue."),
        ("Relative Strength",   "A stock's or sector's performance relative to a benchmark index."),
        ("RNS",                 "Regulatory News Service — the official Stock Exchange channel that all UK listed companies must use to release price-sensitive information."),
        ("ROA",                 "Return on Assets — net income ÷ total assets."),
        ("ROCE",                "Return on Capital Employed — operating profit ÷ capital employed."),
        ("ROE",                 "Return on Equity — net income ÷ shareholders' equity."),
        ("ROIC",                "Return on Invested Capital — net operating profit after tax ÷ invested capital."),
        ("Sector Rotation",     "The movement of investment capital between industry sectors as the economic cycle evolves."),
        ("Ticker",              "A short code identifying a listed stock (e.g. AZN.L for AstraZeneca London)."),
        ("Volatility",          "The degree of price fluctuation. Higher volatility = higher uncertainty = higher risk."),
        ("Yield Curve",         "A graph of government bond yields across different maturities. Its shape signals economic expectations."),
        ("Z-Score (Altman)",    "A statistical measure predicting probability of corporate bankruptcy within 2 years."),
        ("Z-Score (statistical)","How many standard deviations a value is from its historical average."),
    ]
    make_table(doc,
        ["Term", "Definition"],
        glossary,
        col_widths=[4.5, 11.5]
    )
    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Appendix A — Methodology References
# ─────────────────────────────────────────────────────────────────────────────

def add_appendix_a(doc):
    h1(doc, "Appendix A — Methodology References")
    body(doc,
        "The quantitative indicators used in Alpha Move AI are grounded in peer-reviewed academic "
        "research. Below are the key papers underlying each methodology."
    )

    refs = [
        {
            "title": "Price Momentum",
            "papers": [
                {
                    "citation": "Jegadeesh, N. & Titman, S. (1993). 'Returns to Buying Winners and Selling Losers: Implications for Stock Market Efficiency.' Journal of Finance, 48(1), pp. 65–91.",
                    "summary": "The seminal paper demonstrating that stocks which perform well over the past 3–12 months continue to outperform over the following 3–12 months. This 'momentum effect' has since been replicated across virtually every developed and emerging market studied. The 12-1 month window (excluding the most recent month to avoid short-term reversal) comes from this research.",
                },
                {
                    "citation": "Asness, C., Moskowitz, T. & Pedersen, L. (2013). 'Value and Momentum Everywhere.' Journal of Finance, 68(3), pp. 929–985.",
                    "summary": "Demonstrates that momentum works across asset classes (stocks, bonds, currencies, commodities) and geographies, and shows the diversification benefit of combining momentum with value.",
                },
            ]
        },
        {
            "title": "Piotroski F-Score",
            "papers": [
                {
                    "citation": "Piotroski, J.D. (2000). 'Value Investing: The Use of Historical Financial Statement Information to Separate Winners from Losers.' Journal of Accounting Research, 38 (Supplement), pp. 1–41.",
                    "summary": "Piotroski showed that among high book-to-market stocks, a simple 9-point scoring system based on improving fundamentals could separate subsequent strong performers from the weak. High-F-Score value stock portfolios significantly outperformed in his study. The score has since been validated across many markets and time periods.",
                },
            ]
        },
        {
            "title": "Altman Z-Score",
            "papers": [
                {
                    "citation": "Altman, E.I. (1968). 'Financial Ratios, Discriminant Analysis and the Prediction of Corporate Bankruptcy.' Journal of Finance, 23(4), pp. 589–609.",
                    "summary": "Altman used multiple discriminant analysis on 66 manufacturing firms to identify the financial ratios most predictive of bankruptcy within two years. The original model had 72% accuracy two years before failure. The five-variable model combining working capital, retained earnings, operating income, market equity, and revenue relative to total assets remains widely used by credit analysts and risk managers.",
                },
                {
                    "citation": "Altman, E.I. (2000). 'Predicting Financial Distress of Companies: Revisiting the Z-Score and ZETA Models.' Working Paper, Stern School of Business, NYU.",
                    "summary": "Updated validation confirming continued predictive power and discussing adaptations for service industries and international markets.",
                },
            ]
        },
        {
            "title": "Quality / Profitability Premium",
            "papers": [
                {
                    "citation": "Novy-Marx, R. (2013). 'The Other Side of Value: The Gross Profitability Premium.' Journal of Financial Economics, 108(1), pp. 1–28.",
                    "summary": "Demonstrates that highly profitable firms (measured by gross profit / assets) generate significantly higher returns than unprofitable firms, even controlling for value factors. Academic basis for rewarding high gross margins in the Quality Score.",
                },
                {
                    "citation": "Fama, E.F. & French, K.R. (2015). 'A Five-Factor Asset Pricing Model.' Journal of Financial Economics, 116(1), pp. 1–22.",
                    "summary": "The expanded Fama-French model adds profitability and investment factors to the classic three-factor model, confirming that high-profitability firms earn a persistent premium.",
                },
            ]
        },
        {
            "title": "Sector Rotation and the Economic Cycle",
            "papers": [
                {
                    "citation": "Stovall, S. (1996). Standard & Poor's Guide to Sector Investing. McGraw-Hill.",
                    "summary": "The foundational practical framework for understanding which sectors tend to lead and lag at each phase of the economic cycle. The rotation pattern described here has been the practitioner's standard since the 1990s.",
                },
            ]
        },
        {
            "title": "Investor Sentiment (Fear & Greed)",
            "papers": [
                {
                    "citation": "Baker, M. & Wurgler, J. (2007). 'Investor Sentiment in the Stock Market.' Journal of Economic Perspectives, 21(2), pp. 129–151.",
                    "summary": "Documents how investor sentiment — measurable through market breadth, new highs/lows, and other indicators — is predictive of subsequent returns. When sentiment is high (greed), subsequent returns tend to be lower; when low (fear), returns tend to be higher. Particularly strong for smaller, harder-to-arbitrage stocks.",
                },
            ]
        },
    ]

    for section in refs:
        h2(doc, section["title"])
        for paper in section["papers"]:
            # Citation in teal mono-ish
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(paper["citation"])
            r.font.name   = FONT_BODY
            r.font.size   = Pt(10)
            r.font.bold   = True
            r.font.color.rgb = TEAL
            # Summary
            body(doc, paper["summary"], indent=True)

    add_page_break_v2(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Appendix B — ICB Sector Company List
# ─────────────────────────────────────────────────────────────────────────────

def add_appendix_b(doc):
    h1(doc, "Appendix B — ICB Sector Company List")
    body(doc,
        "The Industry Classification Benchmark (ICB) is the taxonomy used to classify UK-listed "
        "companies by sector. Below are all companies tracked by Alpha Move AI in each sector."
    )

    sectors = [
        {
            "name": "Energy",
            "note": "Highly cyclical. Profits closely tied to the oil price. Shell and BP are among the largest FTSE 100 companies. Revenue is partly in USD, so GBP/USD movements affect reported UK earnings.",
            "companies": [
                ("SHEL.L", "Shell",           "Integrated oil & gas major; world's largest listed energy company by market cap"),
                ("BP.L",   "BP",              "Integrated oil & gas major; transitioning toward renewable energy"),
                ("HBR.L",  "Harbour Energy",  "UK's largest independent oil & gas producer; North Sea focused"),
            ]
        },
        {
            "name": "Financials",
            "note": "Banks benefit from rising interest rates but face credit risk in recessions. Insurance companies are sensitive to claims inflation and investment yields. P/B ratio is particularly relevant for banks.",
            "companies": [
                ("HSBA.L", "HSBC",                     "One of the world's largest banks; Asia-focused revenue mix"),
                ("LLOY.L", "Lloyds Banking Group",     "UK's largest retail and commercial bank; highly domestic exposure"),
                ("BARC.L", "Barclays",                 "UK bank with significant investment banking operations"),
                ("NWG.L",  "NatWest Group",             "UK retail and commercial bank; majority government-owned until recent years"),
                ("LSEG.L", "London Stock Exchange Group","Financial data and infrastructure; operates FTSE Russell indices"),
                ("STAN.L", "Standard Chartered",        "Emerging markets-focused international bank"),
                ("AV.L",   "Aviva",                    "UK's largest composite insurer"),
                ("LGEN.L", "Legal & General",           "Major insurer and asset manager"),
                ("ADM.L",  "Admiral Group",             "UK motor and home insurance specialist"),
                ("MNG.L",  "M&G",                      "Investment and savings provider, formerly part of Prudential UK"),
                ("PRU.L",  "Prudential",                "Global insurer with major Asia and Africa operations"),
                ("SDR.L",  "Schroders",                 "Global asset manager with strong UK heritage"),
                ("III.L",  "3i Group",                  "Private equity and infrastructure investment company"),
            ]
        },
        {
            "name": "Industrials",
            "note": "Highly diverse. Defence names like BAE Systems are increasingly seen as defensive given rising government spending. Cyclical industrials like Ashtead are closely tied to construction activity.",
            "companies": [
                ("RR.L",   "Rolls-Royce",      "Jet engine manufacturer and power systems; major recovery story"),
                ("BA.L",   "BAE Systems",      "UK's largest defence contractor; armoured vehicles, aircraft, naval vessels"),
                ("AHT.L",  "Ashtead Group",    "Equipment rental (Sunbelt brand in North America); highly cyclical"),
                ("IAG.L",  "Int'l Airlines Group", "Owner of British Airways, Iberia, Vueling; aviation holding company"),
                ("IMI.L",  "IMI",              "Precision engineering for industrial fluid and motion control"),
                ("WEIR.L", "Weir Group",       "Mining equipment and services"),
                ("RTO.L",  "Rentokil Initial", "Pest control and hygiene services; global scale"),
                ("ITRK.L", "Intertek Group",   "Testing, inspection and certification services"),
                ("MRO.L",  "Melrose Industries","Aerospace components and engineering services"),
                ("EXPN.L", "Experian",         "Global data analytics and credit reporting"),
                ("WPP.L",  "WPP",             "World's largest advertising and marketing services group"),
                ("PSON.L", "Pearson",          "Educational publishing and digital learning"),
            ]
        },
        {
            "name": "Materials",
            "note": "Highly cyclical. Profits driven by commodity prices (iron ore, copper, gold, silver). CRDA and MNDI are specialty chemicals/packaging — less cyclical than pure miners.",
            "companies": [
                ("RIO.L",  "Rio Tinto",      "Global mining giant; iron ore, copper, aluminium, lithium"),
                ("GLEN.L", "Glencore",       "Diversified mining and commodity trading conglomerate"),
                ("AAL.L",  "Anglo American", "Diversified mining; diamonds (De Beers), copper, platinum"),
                ("ANTO.L", "Antofagasta",    "Chilean copper miner; major London-listed copper play"),
                ("FRES.L", "Fresnillo",      "World's largest primary silver producer; also gold mines in Mexico"),
                ("MNDI.L", "Mondi",          "Sustainable packaging and paper"),
                ("SKG.L",  "Smurfit WestRock","Global leader in corrugated packaging"),
                ("CRDA.L", "Croda International","Specialty chemicals for beauty, health and crop protection"),
            ]
        },
        {
            "name": "Consumer Discretionary",
            "note": "Sensitive to consumer confidence and disposable income. Performs well in economic expansions; suffers in downturns. Housebuilders are highly sensitive to mortgage rates.",
            "companies": [
                ("CPG.L",  "Compass Group",    "World's largest contract catering company"),
                ("NXT.L",  "Next",             "UK fashion and homeware retailer; strong online presence"),
                ("IHG.L",  "IHG Hotels & Resorts","Owner of Holiday Inn, Crowne Plaza, InterContinental brands"),
                ("GAW.L",  "Games Workshop",   "Warhammer tabletop games; exceptional brand loyalty and margins"),
                ("KGF.L",  "Kingfisher",       "DIY retail (B&Q, Screwfix, Castorama)"),
                ("JD.L",   "JD Sports Fashion","Sports fashion retailer with global expansion"),
                ("MKS.L",  "Marks & Spencer",  "UK clothing and food retailer; ongoing business transformation"),
                ("WTB.L",  "Whitbread",        "Owner of Premier Inn hotels and restaurants"),
                ("EZJ.L",  "easyJet",          "European low-cost airline"),
                ("ENT.L",  "Entain",           "Sports betting and gaming (Ladbrokes, Coral, bwin)"),
                ("FLTR.L", "Flutter Entertainment","Online gambling (Paddy Power, Betfair, FanDuel)"),
                ("ABF.L",  "Associated British Foods","Primark fashion, plus food ingredients (AB Foods)"),
                ("SBRY.L", "J Sainsbury",      "UK supermarket chain; also Argos and Habitat brands"),
                ("PSN.L",  "Persimmon",        "UK housebuilder; volume builder of affordable homes"),
                ("TW.L",   "Taylor Wimpey",    "UK housebuilder with major residential developments"),
            ]
        },
        {
            "name": "Consumer Staples",
            "note": "Defensive sector. Demand is stable regardless of the economic cycle. Strong brands and pricing power typically lead to high, durable gross margins. Key safe-haven sector during downturns.",
            "companies": [
                ("BATS.L", "British American Tobacco","Cigarettes, vaping, oral nicotine (Lucky Strike, Dunhill, Vuse)"),
                ("ULVR.L", "Unilever",               "Consumer goods giant (Dove, Hellmann's, Magnum, Persil)"),
                ("RKT.L",  "Reckitt Benckiser",       "Health and hygiene brands (Dettol, Nurofen, Durex)"),
                ("TSCO.L", "Tesco",                   "UK's largest supermarket; also Tesco Bank"),
                ("DGE.L",  "Diageo",                  "World's largest spirits company (Johnnie Walker, Guinness, Tanqueray)"),
                ("IMB.L",  "Imperial Brands",         "Tobacco company (Davidoff, West, blu vaping)"),
            ]
        },
        {
            "name": "Health Care",
            "note": "Defensive, driven by healthcare demand rather than economic cycles. AstraZeneca is the FTSE 100's largest company. Drug development creates binary risk events (trial success/failure).",
            "companies": [
                ("AZN.L", "AstraZeneca",         "Global pharmaceutical and oncology leader; largest FTSE 100 company"),
                ("GSK.L", "GSK",                 "Pharmaceuticals and vaccines (formerly GlaxoSmithKline)"),
                ("HLN.L", "Haleon",              "Consumer health brands (Sensodyne, Voltaren, Panadol); GSK spinoff"),
                ("SN.L",  "Smith & Nephew",      "Medical devices (knee and hip implants, wound care)"),
                ("HIK.L", "Hikma Pharmaceuticals","Generic and branded pharmaceuticals; US and MENA markets"),
            ]
        },
        {
            "name": "Technology",
            "note": "The UK tech sector is smaller than the US equivalent. These companies tend to have high gross margins, recurring revenue, and low capital requirements — qualities that drive high Quality Scores.",
            "companies": [
                ("REL.L",  "RELX",          "Information analytics and decision tools (LexisNexis, Elsevier)"),
                ("HLMA.L", "Halma",         "Safety, environmental and healthcare technology products"),
                ("SGE.L",  "Sage Group",    "Accounting and business software for SMEs"),
                ("AUTO.L", "Auto Trader",   "UK's largest digital automotive marketplace"),
                ("RMV.L",  "Rightmove",     "UK's dominant online property portal"),
            ]
        },
        {
            "name": "Telecommunications",
            "note": "Capital-intensive (network infrastructure). Highly regulated. BT and Vodafone have both cut dividends in recent years following heavy investment cycles.",
            "companies": [
                ("VOD.L",   "Vodafone",     "Global telecoms operator; major European and African presence"),
                ("BT-A.L",  "BT Group",     "UK's largest telecoms provider; Openreach fibre network"),
                ("AAF.L",   "Airtel Africa","Mobile telecoms in Africa; subsidiary of Bharti Airtel"),
            ]
        },
        {
            "name": "Utilities",
            "note": "Regulated businesses with predictable, defensive earnings. Sensitive to interest rate changes as investors compare yields to gilt yields. Heavily involved in the UK energy transition.",
            "companies": [
                ("NG.L",  "National Grid",   "UK electricity and gas transmission networks; also US operations"),
                ("SSE.L", "SSE",             "Energy networks and renewable generation (wind, hydro)"),
                ("CNA.L", "Centrica",        "British Gas owner; energy supply and services"),
                ("SVT.L", "Severn Trent",    "Water and waste water services in the Midlands"),
                ("UU.L",  "United Utilities","Water supply and wastewater in North West England"),
            ]
        },
        {
            "name": "Real Estate",
            "note": "REITs must distribute 90% of taxable income as dividends. Net Asset Value (NAV) is the key valuation metric; P/B is particularly relevant. Logistics-focused REITs have benefited from the shift to e-commerce.",
            "companies": [
                ("LAND.L", "Land Securities",       "Commercial property developer and investor (offices, retail)"),
                ("SGRO.L", "Segro",                 "Industrial and logistics real estate; beneficiary of e-commerce growth"),
                ("BLND.L", "British Land",          "Retail parks, offices and mixed-use developments"),
                ("BBOX.L", "Tritax Big Box REIT",   "Large logistics warehouses leased to major retailers"),
                ("PCTN.L", "Picton Property Income","Diversified commercial property investment trust"),
                ("GPE.L",  "Great Portland Estates","Central London offices and mixed-use property"),
            ]
        },
    ]

    for sector in sectors:
        h2(doc, sector["name"])
        callout_box(doc, sector["note"], label="SECTOR NOTE")
        make_table(doc,
            ["Ticker", "Company", "Description"],
            [[c[0], c[1], c[2]] for c in sector["companies"]],
            col_widths=[2.5, 4.5, 9]
        )


# ─────────────────────────────────────────────────────────────────────────────
# Main build function
# ─────────────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page setup — A4, narrow margins for a professional look
    for section in doc.sections:
        section.page_width  = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    add_cover(doc)
    add_toc(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)              # RNS News Screener
    add_section_9_analytics(doc)    # Analytics — Visual Map
    add_section_10_sidebar(doc)     # Sidebar (was Section 8)
    add_section_11_workflows(doc)   # Investment Leads (was Section 9)
    add_section_12_glossary(doc)    # Glossary (was Section 10)
    add_appendix_a(doc)
    add_appendix_b(doc)

    out_path = "Alpha_Move_AI_User_Manual.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_document()
