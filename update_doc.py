"""
Update the Alpha Move AI User Manual Word document with content from the Markdown version.
This script adds missing content and enhances existing sections.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy

SRC = r"c:\Users\richa\Documents\WebProjects\UK_stocks\stock_screener\Alpha_Move_AI_User_Manual.docx"
doc = Document(SRC)


def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = deepcopy(paragraph._element)
    # Clear the new paragraph
    for child in list(new_p):
        new_p.remove(child)
    # Set text
    run_elem = new_p.makeelement(qn("w:r"), {})
    t_elem = run_elem.makeelement(qn("w:t"), {})
    t_elem.text = text
    run_elem.append(t_elem)
    new_p.append(run_elem)
    # Insert after
    paragraph._element.addnext(new_p)
    # Return a Paragraph wrapper
    return Document().add_paragraph(text)  # placeholder, we'll use the element directly


def add_paragraph_after(doc, after_idx, text, bold=False, italic=False):
    """Add a new paragraph after the given paragraph index using XML manipulation."""
    ref_para = doc.paragraphs[after_idx]
    # Create a new paragraph element
    new_p_elem = deepcopy(ref_para._element)
    # Clear it
    for child in list(new_p_elem):
        new_p_elem.remove(child)

    # Create a run with the text
    r_elem = new_p_elem.makeelement(qn("w:r"), {})
    # Add bold if needed
    if bold:
        b_elem = r_elem.makeelement(qn("w:b"), {})
        r_elem.append(b_elem)
    # Add italic if needed
    if italic:
        i_elem = r_elem.makeelement(qn("w:i"), {})
        r_elem.append(i_elem)

    t_elem = r_elem.makeelement(qn("w:t"), {})
    t_elem.text = text
    r_elem.append(t_elem)
    new_p_elem.append(r_elem)

    # Insert after the reference paragraph
    ref_para._element.addnext(new_p_elem)
    return new_p_elem


def add_table_after(doc, after_idx, headers, rows):
    """Add a table after the given paragraph index."""
    ref_para = doc.paragraphs[after_idx]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Set header
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    # Set data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    # Move table after ref_para
    ref_para._element.addnext(table._tbl)
    return table


# ============================================================
# Find key paragraph indices
# ============================================================
targets = {}
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if "Risk Score" in txt and "(1-10)" in txt:
        targets["risk_score_heading"] = i
    elif "A composite assessment" in txt:
        targets["risk_score_desc"] = i
    elif "The Altman Z-Score was developed" in txt:
        targets["altman_desc"] = i
    elif "Save and monitor" in txt and "Watchlist" in txt:
        targets["save_monitor"] = i
    elif "The signals that should all be GREEN" in txt:
        targets["signals_green"] = i
    elif "When eight or nine" in txt:
        targets["eight_nine"] = i

print("Targets found:", targets)

# ============================================================
# UPDATE 1: Enhance Risk Score section
# ============================================================
if "risk_score_desc" in targets:
    idx = targets["risk_score_desc"]

    # Add the detailed breakdown text
    add_paragraph_after(
        doc,
        idx,
        "The score blends two components (60% Altman Z-Score, 40% Volatility):",
        bold=True,
    )

    # Add Altman Z-Score detail
    add_paragraph_after(doc, idx + 1, "Altman Z-Score component:")

    add_paragraph_after(
        doc,
        idx + 2,
        "The Altman Z-Score was developed in 1968 to predict corporate bankruptcy. "
        "It combines five financial ratios into a single number. A Z-Score above 3.0 indicates "
        "a financially safe company; below 1.8 suggests significant distress risk. "
        "The tool converts this to the 1-10 scale:",
    )

    add_paragraph_after(doc, idx + 3, "Z \u2265 3.0 \u2192 Risk score 1 (safe)")
    add_paragraph_after(doc, idx + 4, "Z \u2264 1.0 \u2192 Risk score 10 (distress)")

    add_paragraph_after(doc, idx + 5, "Volatility component:", bold=True)

    add_paragraph_after(
        doc,
        idx + 6,
        "Annualised price volatility (standard deviation of daily log returns \u00d7 \u221a252). "
        "Converted to 1-10:",
    )

    add_paragraph_after(doc, idx + 7, "< 10% annualised \u2192 score 1 (very low)")
    add_paragraph_after(doc, idx + 8, "Each additional 5% adds roughly 1 point")
    add_paragraph_after(doc, idx + 9, "> 60% \u2192 score 10 (very high)")

    print("UPDATE 1: Risk Score section enhanced")

# ============================================================
# UPDATE 2: Add Cycle Phase Suggestion to Fear & Greed
# ============================================================
fg_score_table_idx = None
for i, p in enumerate(doc.paragraphs):
    if "0-24" in p.text and "Extreme Fear" in p.text:
        fg_score_table_idx = i
        break

if fg_score_table_idx:
    add_paragraph_after(doc, fg_score_table_idx, "")
    add_paragraph_after(
        doc,
        fg_score_table_idx + 1,
        "Cycle Phase Suggestion: The tool automatically suggests a market cycle phase "
        "(Recovery / Expansion / Slowdown / Contraction) based on the F&G score and its trend over time.",
        italic=True,
    )
    print("UPDATE 2: Cycle Phase Suggestion added to Fear & Greed")

# ============================================================
# UPDATE 3: Add the "signals that should all be GREEN" table to Section 11.8
# ============================================================
if "save_monitor" in targets:
    idx = targets["save_monitor"]

    # Add the signals table header
    add_paragraph_after(doc, idx, "")
    add_paragraph_after(
        doc, idx + 1, "The signals that should all be GREEN before you buy:", bold=True
    )

    # Add the signals table
    signals_headers = ["Signal", "Source", "Threshold for GREEN"]
    signals_rows = [
        ["Macro backdrop", "Sidebar Fear & Greed", "< 75"],
        ["Sector leadership", "Rotation page", "Sector RS > 1.05"],
        ["Business quality", "Screener", "Quality \u2265 7, Piotroski \u2265 6"],
        ["Valuation", "Screener", "PEGY \u2264 1.5, P/E sector-comparable"],
        ["Trend", "Screener / Chart", "Momentum \u2265 6, price above MA50"],
        ["Safety", "Screener / Health tab", "Risk \u2264 5, Net Debt manageable"],
        [
            "Professional view",
            "Analyst Monitor",
            "Buy consensus, Upside \u2265 10%, Rev Score > 0",
        ],
        [
            "Catalysts",
            "RNS News, Company News",
            "No negative Tier A in 30 days; ideally a recent positive one",
        ],
        [
            "Visual confirmation",
            "Analytics",
            "Green quadrant in both Quality \u00d7 PEGY and Momentum \u00d7 Risk",
        ],
    ]

    # Add the table after the header
    ref_para = doc.paragraphs[idx + 2]
    table = doc.add_table(rows=1 + len(signals_rows), cols=3)
    for i, h in enumerate(signals_headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(signals_rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    ref_para._element.addnext(table._tbl)

    # Add the summary text
    add_paragraph_after(
        doc,
        idx + 3,
        "When eight or nine of these are green for a single name, you have a high-conviction, "
        "asymmetric-risk setup. When fewer than five are green, walk away \u2014 there will be "
        "better opportunities.",
    )

    print("UPDATE 3: Signals table added to Section 11.8")

# ============================================================
# UPDATE 4: Enhance Section 10 - Sidebar with ICB Sectors detail
# ============================================================
sidebar_table = None
for t in doc.tables:
    if t.rows[0].cells[0].text.strip() == "Item" and len(t.rows) >= 7:
        sidebar_table = t
        break

if sidebar_table:
    has_icb = False
    for row in sidebar_table.rows:
        if "ICB" in row.cells[0].text:
            has_icb = True
            break

    if not has_icb:
        row = sidebar_table.add_row()
        row.cells[0].text = "ICB Sectors"
        row.cells[1].text = (
            "Daily percentage change for each of the 11 ICB sectors, calculated as the average return of representative stocks in each sector"
        )
        print("UPDATE 4: ICB Sectors row added to Sidebar table")

# ============================================================
# UPDATE 5: Add the "New Highs vs Lows" component to Fear & Greed
# ============================================================
fg_components_table = None
for t in doc.tables:
    if t.rows[0].cells[0].text.strip() == "Component" and len(t.rows) >= 6:
        fg_components_table = t
        break

if fg_components_table:
    has_hl = False
    for row in fg_components_table.rows:
        if "New Highs" in row.cells[0].text:
            has_hl = True
            break

    if not has_hl:
        row = fg_components_table.add_row()
        row.cells[0].text = "New Highs vs Lows"
        row.cells[1].text = (
            "Ratio of stocks hitting 52-week highs versus 52-week lows across the FTSE universe. More highs = greed; more lows = fear."
        )
        print("UPDATE 5: New Highs vs Lows added to Fear & Greed components")

# ============================================================
# SAVE
# ============================================================
OUTPUT = r"c:\Users\richa\Documents\WebProjects\UK_stocks\stock_screener\Alpha_Move_AI_User_Manual.docx"
doc.save(OUTPUT)
print(f"\nDocument saved to: {OUTPUT}")
print("Update complete!")
