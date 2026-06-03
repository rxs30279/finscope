"""Build Alpha_Move_AI_User_Manual.docx from USER_MANUAL.md.

USER_MANUAL.md is the single source of truth for the user manual. This script
renders it to a styled Word document so the two never drift apart again. It
replaces the old hand-written patch script (update_doc.py), which edited the
.docx directly and is what let it fall out of sync.

What it produces:
  * A cover page with the existing cover image (manual_assets/cover.png).
  * A pre-built, clickable table of contents (static internal hyperlinks, so it
    renders complete in every viewer with no field update / F9 required).
  * Word built-in Heading 1-4 styles for the markdown headings, so the TOC and
    the navigation pane work.
  * GFM pipe tables, bullet/numbered lists, fenced code blocks, inline
    bold/italic/code, and horizontal rules.
  * The "Under the hood" blockquotes rendered as shaded, bordered callout boxes.

Usage:
    python build_manual.py            # MD at repo root -> docx at repo root
    python build_manual.py in.md out.docx

Requires python-docx (already in the toolchain): pip install python-docx
"""

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_ROOT = os.path.dirname(os.path.abspath(__file__))
_DEFAULT_MD = os.path.join(_ROOT, "USER_MANUAL.md")
_DEFAULT_DOCX = os.path.join(_ROOT, "Alpha_Move_AI_User_Manual.docx")
_COVER = os.path.join(_ROOT, "manual_assets", "cover.png")

CODE_FONT = "Consolas"
CALLOUT_FILL = "F2F4F7"   # light grey-blue for the "Under the hood" boxes
HEADER_FILL = "E7ECF3"    # table header shading
CODE_FILL = "F5F5F5"
RULE_COLOR = "CCCCCC"

# Counts top-level headings as they are rendered, so each gets a stable bookmark
# (mh0, mh1, …) that the static table of contents can link to. Reset per build.
_heading_counter = [0]

# ── Markdown parsing ────────────────────────────────────────────────────────
# The manual is hand-written, so we only support the constructs it actually uses
# rather than a full CommonMark parser.

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^(\s*)[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
_TABLE_SEP = re.compile(r"^\s*\|?[\s:\-|]+\|?\s*$")


def parse_blocks(lines):
    """Parse a list of raw markdown lines into a flat list of (type, payload) blocks.

    Block types: heading, paragraph, bullets, numbered, table, code, quote, hr.
    Blockquotes are parsed recursively (their payload is a list of inner blocks).
    """
    blocks = []
    i, n = 0, len(lines)
    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            blocks.append(("hr", None))
            i += 1
            continue

        m = _HEADING.match(raw)
        if m:
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        if stripped.startswith("```"):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append(("code", code))
            continue

        if stripped.startswith(">"):
            inner = []
            while i < n and lines[i].lstrip().startswith(">"):
                t = lines[i].lstrip()[1:]
                if t.startswith(" "):
                    t = t[1:]
                inner.append(t)
                i += 1
            blocks.append(("quote", parse_blocks(inner)))
            continue

        # Table: a pipe row immediately followed by a separator row.
        if "|" in raw and i + 1 < n and _TABLE_SEP.match(lines[i + 1]) and "-" in lines[i + 1]:
            header = _split_row(raw)
            i += 2  # consume header + separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append(("table", (header, rows)))
            continue

        if _BULLET.match(raw) or _NUMBERED.match(raw):
            items = []
            while i < n:
                bm = _BULLET.match(lines[i])
                nm = _NUMBERED.match(lines[i])
                if bm:
                    level = len(bm.group(1)) // 2
                    items.append((level, None, bm.group(2).strip()))
                    i += 1
                elif nm:
                    level = len(nm.group(1)) // 2
                    items.append((level, nm.group(2), nm.group(3).strip()))
                    i += 1
                else:
                    break
            blocks.append(("list", items))
            continue

        # Otherwise: a paragraph runs until a blank line or a new block starter.
        para = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            ns = nxt.strip()
            if ns == "" or ns == "---" or nxt.lstrip().startswith(">"):
                break
            if _HEADING.match(nxt) or ns.startswith("```"):
                break
            if _BULLET.match(nxt) or _NUMBERED.match(nxt):
                break
            if "|" in nxt and "|" in lines[i - 1]:  # likely table continuation
                break
            para.append(ns)
            i += 1
        blocks.append(("paragraph", " ".join(para)))
    return blocks


def _split_row(line):
    """Split a GFM table row into trimmed cell strings."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


# ── Inline formatting ───────────────────────────────────────────────────────
_INLINE = re.compile(
    r"(\*\*.+?\*\*"          # bold
    r"|\*[^*]+?\*"           # italic
    r"|`[^`]+?`"             # inline code
    r"|\[[^\]]+?\]\([^)]+?\))"  # link
)


def add_runs(paragraph, text):
    """Add text to a paragraph, honouring **bold**, *italic*, `code`, and links."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = CODE_FONT
            r.font.color.rgb = RGBColor(0xB1, 0x1A, 0x4B)
        elif tok.startswith("*"):
            paragraph.add_run(tok[1:-1]).italic = True
        else:  # link — render the visible label only
            lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            paragraph.add_run(lm.group(1))
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ── Low-level docx helpers ──────────────────────────────────────────────────

def _shade(element, fill):
    """Apply a background shade to a <w:tcPr> or <w:pPr> element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def add_hr(container):
    p = container.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE_COLOR)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_bookmark(paragraph, name, bid):
    """Wrap a paragraph's content in a Word bookmark so it can be linked to."""
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    p_pr = p.find(qn("w:pPr"))
    if p_pr is not None:
        p_pr.addnext(start)
    else:
        p.insert(0, start)
    p.append(end)


def add_internal_link(paragraph, anchor, text):
    """Add a clickable internal hyperlink (to a bookmark) — renders without F9."""
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def render_static_toc(container, entries):
    """Render a pre-built, clickable table of contents (no field, no F9 needed).

    *entries* is a list of (level, text, anchor). Page numbers are intentionally
    omitted — they cannot be computed without a layout engine — but every entry
    is a working internal hyperlink that jumps to its heading.
    """
    for level, text, anchor in entries:
        p = container.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
        p.paragraph_format.space_after = Pt(2)
        add_internal_link(p, anchor, text)


# ── Block rendering ─────────────────────────────────────────────────────────

def render_table(container, header, rows, *, nested=False):
    table = container.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for ci, text in enumerate(header):
        hdr[ci].paragraphs[0].text = ""
        run = hdr[ci].paragraphs[0].add_run()
        add_runs(hdr[ci].paragraphs[0], text)
        for r in hdr[ci].paragraphs[0].runs:
            r.bold = True
        shade_cell(hdr[ci], HEADER_FILL)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci in range(len(header)):
            cell = cells[ci]
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], row[ci] if ci < len(row) else "")
    return table


def render_list(container, items):
    for level, number, text in items:
        p = container.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
        p.paragraph_format.space_after = Pt(2)
        marker = f"{number}. " if number else ("– " if level else "• ")
        p.add_run(marker)
        add_runs(p, text)


def render_code(container, code_lines):
    p = container.add_paragraph()
    _shade(p._p.get_or_add_pPr(), CODE_FILL)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("\n".join(code_lines))
    run.font.name = CODE_FONT
    run.font.size = Pt(9.5)


def render_quote(doc, inner_blocks):
    """Render a blockquote as a single-cell, shaded, bordered callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade_cell(cell, CALLOUT_FILL)
    first = cell.paragraphs[0]
    render_blocks(cell, inner_blocks)
    # Drop the cell's original empty leading paragraph if we added our own content.
    if not first.text and len(cell.paragraphs) > 1:
        first._element.getparent().remove(first._element)
    # A table cell must not end on a table — guarantee a trailing paragraph.
    if cell._tc.findall(qn("w:tbl")):
        last = cell.paragraphs[-1] if cell.paragraphs else None
        if last is None or last.text or last._element.getprevious() is not None and \
                last._element.getprevious().tag == qn("w:tbl"):
            cell.add_paragraph()


def render_blocks(container, blocks, *, top_level=False, toc_entries=None):
    """Render parsed blocks into a container (Document or table cell)."""
    skip_next_list = False
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            if top_level and text.lower() == "table of contents":
                _page_break(container)
                _add_heading(container, 1, text)
                render_static_toc(container, toc_entries or [])
                # The markdown TOC list immediately follows — skip it.
                skip_next_list = True
                _page_break(container)
                continue
            _add_heading(container, level, text, bookmark=top_level)
        elif kind == "paragraph":
            p = container.add_paragraph()
            add_runs(p, payload)
        elif kind == "list":
            if skip_next_list:
                skip_next_list = False
                continue
            render_list(container, payload)
        elif kind == "table":
            render_table(container, payload[0], payload[1])
        elif kind == "code":
            render_code(container, payload)
        elif kind == "quote":
            render_quote(container, payload)
        elif kind == "hr":
            add_hr(container)


def _add_heading(container, level, text, bookmark=False):
    style = f"Heading {min(level, 4)}"
    p = container.add_paragraph(style=style)
    add_runs(p, text)
    if bookmark:
        idx = _heading_counter[0]
        add_bookmark(p, f"mh{idx}", idx)
        _heading_counter[0] += 1
    return p


def _page_break(container):
    p = container.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def prescan_headings(blocks):
    """Assign each top-level heading a bookmark anchor and collect TOC entries.

    Mirrors the order in which render_blocks assigns bookmarks: every top-level
    heading except "Table of Contents" is bookmarked mh0, mh1, … . Only levels
    2 and 3 (sections and subsections) go into the visible TOC, which keeps the
    title, the TOC page itself, and the #### score-detail headings out of it.
    """
    entries = []
    idx = 0
    for kind, payload in blocks:
        if kind != "heading":
            continue
        level, text = payload
        if text.lower() == "table of contents":
            continue
        anchor = f"mh{idx}"
        idx += 1
        if level in (2, 3):
            entries.append((level, text, anchor))
    return entries


# ── Top-level build ─────────────────────────────────────────────────────────

def build(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    blocks = parse_blocks(lines)
    _heading_counter[0] = 0
    toc_entries = prescan_headings(blocks)

    doc = Document()
    _set_base_style(doc)

    # Cover image, centred, on the first page.
    if os.path.isfile(_COVER):
        cover_p = doc.add_paragraph()
        cover_p.alignment = 1  # centre
        cover_p.add_run().add_picture(_COVER, width=Inches(5.5))
    else:
        print(f"WARNING: cover image not found at {_COVER}; skipping cover.")

    render_blocks(doc, blocks, top_level=True, toc_entries=toc_entries)

    doc.save(docx_path)
    return doc


def _set_base_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)


def main():
    md = sys.argv[1] if len(sys.argv) > 1 else _DEFAULT_MD
    out = sys.argv[2] if len(sys.argv) > 2 else _DEFAULT_DOCX
    if not os.path.isfile(md):
        print(f"Markdown source not found: {md}")
        sys.exit(1)
    build(md, out)
    size = os.path.getsize(out)
    print(f"Built {out} ({size:,} bytes) from {os.path.basename(md)}.")
    print("Table of contents is pre-built and clickable — no field update needed.")


if __name__ == "__main__":
    main()
